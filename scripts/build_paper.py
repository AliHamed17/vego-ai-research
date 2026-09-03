"""Render the paper markdown to a compact DOCX with embedded figures.

Markdown subset: '# ' title, '## ' section, '**Abstract.**'-style bold lead-ins,
paragraphs, '- ' bullets, pipe tables, **bold**/*italic* inline, and a figure
directive line of the form  FIGURE:<file.png>|<caption>.
"""
from __future__ import annotations

import argparse
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Calibri"
BODY_PT = 9.5
SMALL_PT = 8.2
INK = RGBColor(0x0B, 0x0B, 0x0B)
INK2 = RGBColor(0x52, 0x51, 0x4E)
ACCENT = RGBColor(0x17, 0x3B, 0x57)
USABLE_CM = 21.0 - 2 * 1.6


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), fill)
    pr.append(el)


def runs(paragraph, text, size, bold_all=False, color=INK, italic_all=False):
    for tok in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text.replace("`", "")):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(tok)
        run.bold = run.bold or bold_all
        run.italic = run.italic or italic_all
        run.font.size = Pt(size)
        run.font.name = FONT
        run.font.color.rgb = color
        rpr = run._element.get_or_add_rPr()
        rf = rpr.find(qn("w:rFonts"))
        if rf is None:
            rf = OxmlElement("w:rFonts")
            rpr.append(rf)
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(attr), FONT)


def tight(paragraph, before=0.0, after=2.0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_rtl(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:bidi"))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build(md_path, out_path, fig_dir, rtl=False):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, side, Cm(1.6))
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(BODY_PT)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.startswith("FIGURE:"):
            spec = line[len("FIGURE:"):]
            fname, caption = spec.split("|", 1)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tight(para, 4, 1)
            para.add_run().add_picture(os.path.join(fig_dir, fname.strip()), width=Cm(USABLE_CM))
            cap = doc.add_paragraph()
            runs(cap, caption.strip(), SMALL_PT, color=INK2)
            tight(cap, 0, 5)
            if rtl:
                set_rtl(cap)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("# "):
            para = doc.add_paragraph()
            runs(para, line[2:].strip(), 14, bold_all=True, color=ACCENT)
            tight(para, 0, 2)
            if rtl:
                set_rtl(para)
        elif line.startswith("## "):
            para = doc.add_paragraph()
            runs(para, line[3:].strip(), 10.5, bold_all=True, color=ACCENT)
            tight(para, 5, 1.5)
            if rtl:
                set_rtl(para)
        elif line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            ncol = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=ncol)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            comparison_header = [
                "Case",
                "Baseline Issue",
                "Trigger",
                "Human Input",
                "Result After Intervention",
                "Reference",
                "Outcome",
            ]
            if ncol == 7 and rows[0] == comparison_header:
                widths = [1.2, 2.6, 2.5, 2.4, 3.0, 3.0, 3.1]
            else:
                first = 0.22 if ncol >= 4 else 0.3
                widths = [USABLE_CM * first] + [USABLE_CM * (1 - first) / (ncol - 1)] * (ncol - 1)
            for c, width in enumerate(widths):
                table.columns[c].width = Cm(width)
            for r, cells in enumerate(rows):
                for c in range(ncol):
                    cell = table.cell(r, c)
                    cell.width = Cm(widths[c])
                    cell.text = ""
                    para = cell.paragraphs[0]
                    runs(para, cells[c] if c < len(cells) else "", SMALL_PT, bold_all=(r == 0))
                    tight(para, 0.5, 0.5)
                    if rtl:
                        set_rtl(para)
                    if r == 0:
                        shade(cell, "E8EEF3")
            spacer = doc.add_paragraph()
            tight(spacer, 0, 0)
            spacer.paragraph_format.line_spacing = Pt(4)
            continue
        elif re.match(r"^\s*[-*] ", line):
            para = doc.add_paragraph(style="List Bullet")
            runs(para, re.sub(r"^\s*[-*] ", "", line), BODY_PT)
            tight(para, 0, 1)
            para.paragraph_format.left_indent = Cm(0.5)
        else:
            buf = [line.strip()]
            while (i + 1 < len(lines) and lines[i + 1].strip()
                   and not re.match(r"^(#|\||FIGURE:|\s*[-*] )", lines[i + 1])):
                i += 1
                buf.append(lines[i].strip())
            para = doc.add_paragraph()
            runs(para, " ".join(buf), BODY_PT)
            tight(para, 0, 3)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            if rtl:
                set_rtl(para)
        i += 1
    doc.save(out_path)
    print("wrote", out_path)


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("markdown")
    ap.add_argument("docx")
    ap.add_argument("--figures", required=True)
    ap.add_argument("--rtl", action="store_true")
    args = ap.parse_args()
    build(args.markdown, args.docx, args.figures, rtl=args.rtl)
