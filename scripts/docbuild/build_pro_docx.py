#!/usr/bin/env python3
"""Render a Markdown source document into a professionally formatted Word file.

Adds over the plain converter: a title block, a page-numbered footer, shaded
table headers with banded rows, automatic landscape orientation for wide
tables, and consistent typography. Content is never altered - this is
presentation only.

Usage: build_pro_docx.py <source.md> <out.docx> "<Document title>" "<Status line>"
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = Path(sys.argv[1])
OUT = Path(sys.argv[2])
TITLE = sys.argv[3] if len(sys.argv) > 3 else SRC.stem
STATUS = sys.argv[4] if len(sys.argv) > 4 else ""

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
BLUE = RGBColor(0x2F, 0x5A, 0xA8)
GREY = RGBColor(0x5B, 0x64, 0x72)
NAVY_HEX = "1B2A4A"
BAND_HEX = "F5F7FB"

# **bold** must be tried before *italic*, or the bold markers split wrongly.
TOKEN = re.compile(r"\*\*(.+?)\*\*|\*([^*\n]+)\*|`([^`]+)`")


def strip_inline(text: str) -> list[tuple[str, bool, bool, bool]]:
    """Return runs of (text, bold, code, italic)."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    parts: list[tuple[str, bool, bool, bool]] = []
    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False, False))
        if m.group(1) is not None:
            parts.append((m.group(1).replace("`", ""), True, False, False))
        elif m.group(2) is not None:
            parts.append((m.group(2).replace("`", ""), False, False, True))
        else:
            parts.append((m.group(3), False, True, False))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False, False))
    return [p for p in parts if p[0]]


def is_block_start(ln: str) -> bool:
    s = ln.strip()
    return (
        s == "" or s == "---" or s.startswith("#") or s.startswith(">")
        or s.startswith("|") or s.startswith("```") or bool(re.match(r"^[-*]\s+", s))
        or bool(re.match(r"^\d+\.\s+", s))
        # A line opening with a bold label ("**Status:** ...") is its own field
        # line in these documents, not a continuation of the previous sentence.
        or s.startswith("**")
    )


def reflow(raw: list[str]) -> list[str]:
    out: list[str] = []
    i, n = 0, len(raw)
    while i < n:
        ln = raw[i]
        s = ln.strip()
        # Preserve fenced code blocks verbatim, including their line breaks.
        if s.startswith("```"):
            out.append(ln)
            i += 1
            while i < n and not raw[i].strip().startswith("```"):
                out.append(raw[i])
                i += 1
            if i < n:
                out.append(raw[i])
                i += 1
            continue
        if s == "" or s == "---" or s.startswith("#") or s.startswith(">") or s.startswith("|"):
            out.append(ln)
            i += 1
            continue
        buf = ln.rstrip()
        i += 1
        while i < n and not is_block_start(raw[i]):
            buf += " " + raw[i].strip()
            i += 1
        out.append(buf)
    return out


def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_field(paragraph, instr: str) -> None:
    r = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r._r.append(fld)
    r2 = paragraph.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = instr
    r2._r.append(it)
    r3 = paragraph.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r3._r.append(fld2)


lines = SRC.read_text(encoding="utf-8").splitlines()

# Decide orientation from the widest table in the source.
max_cols = 0
for ln in lines:
    if ln.strip().startswith("|"):
        max_cols = max(max_cols, len(ln.strip().strip("|").split("|")))
landscape = max_cols >= 6

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)

sec = doc.sections[0]
if landscape:
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)
else:
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)
for attr, val in (("left_margin", 0.8), ("right_margin", 0.8), ("top_margin", 0.7), ("bottom_margin", 0.7)):
    setattr(sec, attr, Inches(val))

usable_in = (sec.page_width - sec.left_margin - sec.right_margin) / 914400

# ---- title block ----
kicker = doc.add_paragraph()
kr = kicker.add_run("VEGO-AI  ·  PhD RESEARCH PROGRAMME")
kr.font.size = Pt(8)
kr.font.bold = True
kr.font.color.rgb = BLUE
kicker.paragraph_format.space_after = Pt(2)

tp = doc.add_paragraph()
tr = tp.add_run(TITLE)
tr.font.size = Pt(20)
tr.font.bold = True
tr.font.color.rgb = NAVY
tp.paragraph_format.space_after = Pt(4)

if STATUS:
    sp = doc.add_paragraph()
    sr = sp.add_run(STATUS)
    sr.font.size = Pt(9)
    sr.font.italic = True
    sr.font.color.rgb = GREY
    sp.paragraph_format.space_after = Pt(2)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(2)
rule.paragraph_format.space_after = Pt(10)
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "8")
bottom.set(qn("w:color"), NAVY_HEX)
p_bdr.append(bottom)
rule._p.get_or_add_pPr().append(p_bdr)

# ---- footer ----
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(f"{TITLE}   ·   page ")
fr.font.size = Pt(8)
fr.font.color.rgb = GREY
add_field(footer_p, " PAGE ")
fr2 = footer_p.add_run(" of ")
fr2.font.size = Pt(8)
fr2.font.color.rgb = GREY
add_field(footer_p, " NUMPAGES ")
for run in footer_p.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


def add_para(text, *, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    for chunk, bold, code, ital in strip_inline(text):
        run = p.add_run(chunk)
        run.bold = bold
        run.italic = italic or ital
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)
    return p


lines = reflow(lines)
i, n = 0, len(lines)
first_h1_skipped = False

while i < n:
    line = lines[i].rstrip()

    if line.strip().startswith("```"):
        i += 1
        code: list[str] = []
        while i < n and not lines[i].strip().startswith("```"):
            code.append(lines[i].rstrip())
            i += 1
        i += 1  # closing fence
        while code and not code[0].strip():
            code.pop(0)
        while code and not code[-1].strip():
            code.pop()
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Inches(0.2)
        cp.paragraph_format.space_before = Pt(4)
        cp.paragraph_format.space_after = Pt(8)
        for idx, cl in enumerate(code):
            run = cp.add_run(cl)
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            run.font.color.rgb = NAVY
            if idx < len(code) - 1:
                run.add_break()
        shade_para = OxmlElement("w:shd")
        shade_para.set(qn("w:val"), "clear")
        shade_para.set(qn("w:fill"), BAND_HEX)
        cp._p.get_or_add_pPr().append(shade_para)
        continue

    if line.startswith(">"):
        block = []
        while i < n and lines[i].startswith(">"):
            block.append(lines[i].lstrip(">").strip())
            i += 1
        text = " ".join(x for x in block if x)
        p = add_para(text, italic=True, color=GREY, size=9.5)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(8)
        continue

    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m:
        level = len(m.group(1))
        htext = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", m.group(2)).replace("**", "").replace("`", "")
        # The document title is already in the title block; skip a duplicate H1.
        if level == 1 and not first_h1_skipped:
            first_h1_skipped = True
            i += 1
            continue
        h = doc.add_heading(htext, level=min(max(level - 1, 1), 4))
        for r in h.runs:
            r.font.color.rgb = NAVY
            r.font.name = "Calibri"
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        i += 1
        continue

    if line.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|?\s*$", lines[i + 1]):
        header = [c.strip() for c in line.strip("|").split("|")]
        i += 2
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
            i += 1
        ncol = len(header)
        table = doc.add_table(rows=1, cols=ncol)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        col_w = usable_in / ncol
        for c, htext in enumerate(header):
            cell = table.rows[0].cells[c]
            cell.width = Inches(col_w)
            cell.text = ""
            run = cell.paragraphs[0].add_run(re.sub(r"\*\*|`", "", htext))
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade(cell, NAVY_HEX)
        for ri, r in enumerate(rows):
            cells = table.add_row().cells
            for c in range(ncol):
                val = r[c] if c < len(r) else ""
                cells[c].width = Inches(col_w)
                cells[c].text = ""
                para = cells[c].paragraphs[0]
                for chunk, bold, code, ital in strip_inline(val):
                    run = para.add_run(chunk)
                    run.bold = bold
                    run.italic = ital
                    run.font.size = Pt(8.5)
                    if code:
                        run.font.name = "Consolas"
                        run.font.size = Pt(8)
                if ri % 2 == 1:
                    shade(cells[c], BAND_HEX)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        continue

    if re.match(r"^[-*]\s+", line):
        p = doc.add_paragraph(style="List Bullet")
        for chunk, bold, code, ital in strip_inline(re.sub(r"^[-*]\s+", "", line)):
            run = p.add_run(chunk)
            run.bold = bold
            run.italic = ital
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)
        i += 1
        continue

    if re.match(r"^\d+\.\s+", line):
        p = doc.add_paragraph(style="List Number")
        for chunk, bold, code, ital in strip_inline(re.sub(r"^\d+\.\s+", "", line)):
            run = p.add_run(chunk)
            run.bold = bold
            run.italic = ital
            if code:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(3)
        i += 1
        continue

    if line.strip() in ("---", ""):
        i += 1
        continue

    add_para(line)
    i += 1

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"WROTE {OUT.name} ({'landscape' if landscape else 'portrait'}, max_cols={max_cols})")
