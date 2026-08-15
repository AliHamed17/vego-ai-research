#!/usr/bin/env python3
"""Build the evidence-bounded 2026-08-19 supervisor review package."""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import zipfile
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
PACKAGE_ROOT = ROOT / "docs/research/meetings/2026-08-19-supervisor-package"
CONTENT_PATH = PACKAGE_ROOT / "source/package-content.json"
FINAL_DIR = PACKAGE_ROOT / "final"
QA_DIR = PACKAGE_ROOT / "qa"
WORK_DIR = ROOT / "output/aug19-supervisor-package"
RENDER_DIR = WORK_DIR / "render"
TAXONOMY_REPOSITORY_COMMIT = "7b3ba9deefe99172748582f6025d995ccc2a6f86"
SURVEY_URL = "https://aclanthology.org/2026.findings-acl.1811/"
REPOSITORY_URL = (
    "https://github.com/HenryPengZou/Awesome-Human-Agent-Collaboration-Interaction-Systems/"
    f"tree/{TAXONOMY_REPOSITORY_COMMIT}"
)

INK = "17233C"
BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
MUTED = "5B6578"
PANEL = "F4F6F9"
LINE = "C7CFDB"
PURPLE = "7257D6"
WHITE = "FFFFFF"

REQUIRED_FINAL_FILES = (
    "Chapter_2_Literature_Review_EN.docx",
    "Chapter_2_Literature_Review_EN.pdf",
    "Chapter_2_Literature_Review_HE.docx",
    "Chapter_2_Literature_Review_HE.pdf",
    "Human_Agent_Classification_Bilingual.pptx",
    "Human_Agent_Classification_Bilingual.pdf",
    "Supervisor_PreRead_EN.pdf",
    "Supervisor_PreRead_HE.pdf",
    "Supervisor_Tracker_and_Decisions_Bilingual.pdf",
    "package-manifest.sanitized.json",
)

LTR_TOKEN = re.compile(
    r"(?:C2-(?:ACL|KS)-\d{2}|C2-\d{2}|(?:EXP|QL|F|A|D|Q|R)\-?\d+(?:-\d+)?|"
    r"\d+/\d+|Plan [AB]|Section 4|VEGO-AI-PHD-LITERATURE-WORKBOOK|VEGO-AI|"
    r"ACL|URL|DOI|RQ|E6|E8|(?:EN|HE):)"
)
LTR_MARK = "\u200e"
NONBREAKING_HYPHEN = "\u2011"
CANONICAL_CONTROL_TOKEN = re.compile(r"\b(?:F|A|D|Q|R)12-\d{3}\b")


@dataclass(frozen=True)
class BuildRuntime:
    python_exe: Path
    powershell_exe: Path
    node_exe: Path
    node_modules: Path
    override_bin: Path
    docx_sanitizer: Path


def require_runtime_path(path: Path, *, directory: bool = False) -> Path:
    resolved = path.resolve()
    valid = resolved.is_dir() if directory else resolved.is_file()
    if not valid:
        kind = "directory" if directory else "file"
        raise RuntimeError(f"required runtime {kind} is unavailable: {path.name}")
    return resolved


def resolve_pdftoppm(override_bin: Path) -> Path:
    native_executable = (
        override_bin.parent.parent / "native/poppler/Library/bin/pdftoppm.exe"
    )
    if native_executable.is_file():
        return native_executable.resolve()
    for name in ("pdftoppm.cmd", "pdftoppm.exe", "pdftoppm"):
        candidate = override_bin / name
        if candidate.is_file():
            return candidate.resolve()
    raise RuntimeError("bundled pdftoppm wrapper is unavailable")


def load_content() -> dict[str, Any]:
    return json.loads(CONTENT_PATH.read_text(encoding="utf-8"))


def load_private_workbook_binding(
    content: dict[str, Any],
    binding_path: Path | None,
    *,
    ali_approved_release: bool,
) -> tuple[str | None, str | None]:
    """Load an approved live URL only from a private, untracked binding file."""
    if binding_path is None:
        if ali_approved_release:
            raise RuntimeError("Ali-approved release requires a private workbook binding")
        return None, None
    if not ali_approved_release:
        raise RuntimeError("private workbook binding requires explicit Ali-approved release")

    resolved = require_runtime_path(binding_path)
    try:
        relative = resolved.relative_to(ROOT.resolve())
    except ValueError:
        relative = None
    if relative is not None:
        relative_text = relative.as_posix()
        tracked = subprocess.run(
            ["git", "ls-files", "--error-unmatch", "--", relative_text],
            cwd=ROOT,
            capture_output=True,
        )
        ignored = subprocess.run(
            ["git", "check-ignore", "--quiet", "--", relative_text],
            cwd=ROOT,
            capture_output=True,
        )
        if tracked.returncode == 0 or ignored.returncode != 0:
            raise RuntimeError("private workbook binding must be ignored and untracked")

    binding = json.loads(resolved.read_text(encoding="utf-8"))
    expected_alias = str(content.get("workbook", {}).get("logical_alias", ""))
    if binding.get("logical_alias") != expected_alias:
        raise RuntimeError("private workbook binding alias does not match controlled content")
    url = str(binding.get("url", ""))
    if not re.fullmatch(
        r"https://docs\.google\.com/spreadsheets/d/[A-Za-z0-9_-]+(?:/[^\s]*)?",
        url,
    ):
        raise RuntimeError("private workbook binding URL is not an HTTPS Google Sheet")
    return url, sha256_file(resolved)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_text_lf(path: Path, text: str) -> None:
    """Write tracked text with canonical LF bytes on every platform."""
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        handle.write(text)


def run_checked(command: list[str], *, cwd: Path = ROOT) -> None:
    result = subprocess.run(command, cwd=cwd, text=True, capture_output=True, encoding="utf-8")
    if result.returncode:
        raise RuntimeError(
            f"command failed ({result.returncode}): {' '.join(command)}\n"
            f"stdout:\n{result.stdout}\nstderr:\n{result.stderr}"
        )


def set_cell_shading(cell: Any, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell: Any, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    table_width = sum(widths_dxa)
    if table_width not in {9360, 14400}:
        raise ValueError(
            f"table widths must total a controlled portrait or landscape width: {widths_dxa}"
        )
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        set_row_cant_split(row)
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_paragraph_rtl(paragraph: Any) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    # Word resolves physical left/right in the opposite direction for bidi
    # paragraphs.  ``left`` therefore renders Hebrew text at the visual right.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_run_rtl(run: Any, rtl: bool) -> None:
    r_pr = run._r.get_or_add_rPr()
    node = r_pr.find(qn("w:rtl"))
    if node is None:
        node = OxmlElement("w:rtl")
        r_pr.append(node)
    node.set(qn("w:val"), "1" if rtl else "0")


def add_text_runs(paragraph: Any, text: str, *, rtl: bool = False, bold: bool = False) -> None:
    if not rtl:
        run = paragraph.add_run(text)
        run.bold = bold
        return
    set_paragraph_rtl(paragraph)
    cursor = 0
    for match in LTR_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            run.font.name = "Arial"
            run.bold = bold
            set_run_rtl(run, True)
        display_token = match.group(0).replace("-", NONBREAKING_HYPHEN)
        token = paragraph.add_run(f"{LTR_MARK}{display_token}{LTR_MARK}")
        token.font.name = "Arial"
        token.bold = bold
        set_run_rtl(token, False)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.font.name = "Arial"
        run.bold = bold
        set_run_rtl(run, True)


def protect_canonical_control_ids(text: str) -> str:
    """Keep canonical meeting-control IDs intact in non-RTL table cells."""
    return CANONICAL_CONTROL_TOKEN.sub(
        lambda match: match.group(0).replace("-", NONBREAKING_HYPHEN), text
    )


def add_hyperlink(paragraph: Any, label: str, url: str, *, rtl: bool = False) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    run_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rtl_node = OxmlElement("w:rtl")
    rtl_node.set(qn("w:val"), "0" if not rtl else "1")
    run_pr.extend([color, underline, rtl_node])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([run_pr, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_workbook_reference(
    paragraph: Any,
    content: dict[str, Any],
    *,
    rtl: bool,
    bilingual: bool = False,
) -> None:
    workbook = content["workbook"]
    alias = workbook["logical_alias"]
    label = f"{workbook['title']} [{alias}]"
    add_text_runs(paragraph, label, rtl=rtl)
    private_url = content.get("_private_workbook_url")
    if private_url:
        add_text_runs(paragraph, " | ", rtl=rtl)
        add_hyperlink(
            paragraph,
            "Open approved workbook" if not rtl else "פתיחת החוברת המאושרת",
            str(private_url),
            rtl=rtl,
        )
        return
    if bilingual:
        notice = " | Live link withheld pending Ali-approved release | הקישור החי מושהה עד לשחרור באישור עלי"
    elif rtl:
        notice = " | הקישור החי אינו נכלל עד לשחרור באישור עלי"
    else:
        notice = " | Live link withheld pending Ali-approved release"
    add_text_runs(paragraph, notice, rtl=rtl)


def configure_document(doc: Document, *, rtl: bool, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial" if rtl else "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DEEP_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Arial" if rtl else "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    if rtl:
        for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
            style = styles[style_name]
            p_pr = style.element.get_or_add_pPr()
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            p_pr.append(bidi)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    add_text_runs(header, running_label, rtl=rtl, bold=True)
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text_runs(
        footer,
        "VEGO-AI | בדיקת מנחים | 19 באוגוסט 2026"
        if rtl
        else "VEGO-AI | Supervisor review | 19 August 2026",
        rtl=rtl,
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)

    props = doc.core_properties
    props.author = "Ali Hamed"
    props.subject = "Evidence-bounded literature review for supervisor discussion"
    props.keywords = "VEGO-AI, literature review, human-agent collaboration, evidence boundary"


def add_title_block(
    doc: Document,
    title: str,
    subtitle: str,
    *,
    rtl: bool,
    secondary_title: str | None = None,
) -> None:
    kicker = doc.add_paragraph()
    add_text_runs(
        kicker,
        "VEGO-AI | PHD RESEARCH PROGRAMME" if not rtl else "VEGO-AI | תכנית מחקר לדוקטורט",
        rtl=rtl,
        bold=True,
    )
    if not rtl:
        kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in kicker.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(PURPLE)
    kicker.paragraph_format.space_after = Pt(8)
    title_p = doc.add_paragraph()
    add_text_runs(title_p, title, rtl=rtl, bold=True)
    if not rtl:
        title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title_p.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor.from_string(INK)
    title_p.paragraph_format.space_after = Pt(0 if secondary_title else 6)
    if secondary_title:
        secondary_p = doc.add_paragraph()
        add_text_runs(secondary_p, secondary_title, rtl=True, bold=True)
        for run in secondary_p.runs:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor.from_string(INK)
        secondary_p.paragraph_format.space_after = Pt(6)
    subtitle_p = doc.add_paragraph()
    add_text_runs(subtitle_p, subtitle, rtl=rtl)
    if not rtl:
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in subtitle_p.runs:
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle_p.paragraph_format.space_after = Pt(12)


def add_heading(doc: Document, text: str, *, level: int, rtl: bool) -> Any:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    add_text_runs(paragraph, text, rtl=rtl)
    return paragraph


def add_body(doc: Document, text: str, *, rtl: bool) -> Any:
    paragraph = doc.add_paragraph()
    add_text_runs(paragraph, text, rtl=rtl)
    if rtl:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_bullet(doc: Document, text: str, *, rtl: bool, number: int | None = None) -> Any:
    paragraph = doc.add_paragraph()
    marker = f"{number}. " if number is not None else "• "
    add_text_runs(paragraph, marker + text, rtl=rtl)
    if rtl:
        paragraph.paragraph_format.right_indent = Inches(0.375)
    else:
        paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    return paragraph


def add_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[str]],
    *,
    rtl: bool,
    widths_dxa: list[int] | None = None,
    font_size: float = 9,
) -> Any:
    row_values = [list(map(str, row)) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = widths_dxa or {
        2: [3000, 6360],
        3: [2200, 3000, 4160],
        4: [1900, 2400, 1900, 3160],
    }[len(headers)]
    if len(widths) != len(headers):
        raise ValueError("table width count must match the header count")
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, PANEL)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_text_runs(paragraph, header, rtl=rtl, bold=True)
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
    for values in row_values:
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            for part_index, part in enumerate(value.split("\n")):
                paragraph = cell.paragraphs[0] if part_index == 0 else cell.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                part_rtl = rtl or (part.lstrip().startswith("HE:") and bool(re.search(r"[\u0590-\u05ff]", part)))
                display_part = part if part_rtl else protect_canonical_control_ids(part)
                add_text_runs(paragraph, display_part, rtl=part_rtl)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_evidence_callout(doc: Document, lines: list[str], *, rtl: bool) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF2FF")
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3 if index < len(lines) - 1 else 0)
        add_text_runs(paragraph, line, rtl=rtl, bold=(index == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_chapter2(content: dict[str, Any], language: str, output: Path) -> None:
    rtl = language == "he"
    doc = Document()
    configure_document(
        doc,
        rtl=rtl,
        running_label="פרק 2 | ספרות בלבד | טיוטה לבדיקת המנחים"
        if rtl
        else "Chapter 2 | Literature only | Supervisor-review draft",
    )
    doc.core_properties.title = content["chapter2"][f"title_{language}"]
    add_title_block(
        doc,
        content["chapter2"][f"title_{language}"],
        content["chapter2"][f"subtitle_{language}"],
        rtl=rtl,
    )
    add_evidence_callout(
        doc,
        [
            "גבול הראיות: 525 מופעים המתכנסים ל־116 עבודות; סינון אנושי 0/116; EXP-005 0/24; מוכנות רפואית 0/6."
            if rtl
            else "Evidence boundary: 525 occurrences -> 116 works; human screening 0/116; EXP-005 0/24; medical readiness 0/6.",
            "מצב: מוכן לבדיקת עלי; לא נמסר, לא התקבל ולא נסגר."
            if rtl
            else "Status: Ready for Ali review; not delivered, accepted, or closed.",
        ],
        rtl=rtl,
    )
    for section in content["chapter2"][language]:
        add_heading(doc, f"{section['id']}  {section['title']}", level=1, rtl=rtl)
        for paragraph in section.get("paragraphs", []):
            add_body(doc, paragraph, rtl=rtl)
        if section.get("table"):
            table_data = section["table"]
            add_table(doc, table_data["headers"], table_data["rows"], rtl=rtl)
        if section.get("citations"):
            label = "עוגני מקור: " if rtl else "Source anchors: "
            p = doc.add_paragraph()
            add_text_runs(p, label + "; ".join(section["citations"]), rtl=rtl)
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor.from_string(MUTED)

    add_heading(doc, "מקורות ועוגני ראיות" if rtl else "Sources and evidence anchors", level=1, rtl=rtl)
    intro = (
        "עוגני ACL להלן נבדקו ברמת עמוד. שתים-עשרה העבודות הנוספות הן מועמדות לסינון אנושי בלבד."
        if rtl
        else "The ACL anchors below were reviewed at page level. The twelve additional works remain human-screening candidates only."
    )
    add_body(doc, intro, rtl=rtl)
    for row in content["source_anchors"]:
        p = doc.add_paragraph()
        add_text_runs(
            p,
            f"{row['id']} | {row['locator']} | "
            + (row["safe_statement_he"] if rtl else row["safe_statement_en"]),
            rtl=rtl,
            bold=True,
        )
        p2 = doc.add_paragraph()
        add_text_runs(
            p2,
            ("גבול: " if rtl else "Boundary: ")
            + (row["boundary_he"] if rtl else row["boundary_en"])
            + " ",
            rtl=rtl,
        )
        add_hyperlink(p2, SURVEY_URL, SURVEY_URL)
    add_heading(doc, "מועמדות שמרניות לסינון" if rtl else "Conservative screening candidates", level=2, rtl=rtl)
    for row in content["candidate_sources"]:
        p = doc.add_paragraph()
        add_text_runs(
            p,
            f"{row['id']} | {row['title']} | " + (row["use_he"] if rtl else row["use_en"]) + " | ",
            rtl=rtl,
        )
        add_hyperlink(p, row["url"], row["url"])
    p = doc.add_paragraph()
    add_text_runs(p, "חוברת ספרות מקורית: " if rtl else "Native literature workbook: ", rtl=rtl)
    add_workbook_reference(p, content, rtl=rtl)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_preread(content: dict[str, Any], language: str, output: Path) -> None:
    rtl = language == "he"
    doc = Document()
    configure_document(
        doc,
        rtl=rtl,
        running_label="קריאה מקדימה | 19 באוגוסט 2026" if rtl else "Supervisor pre-read | 19 August 2026",
    )
    title = "קריאה מקדימה למנחים - ספרות בלבד" if rtl else "Supervisor pre-read - literature only"
    subtitle = (
        "19 באוגוסט 2026 | 09:00-10:00 | Asia/Jerusalem"
        if rtl
        else "19 August 2026 | 09:00-10:00 | Asia/Jerusalem"
    )
    add_title_block(doc, title, subtitle, rtl=rtl)
    add_evidence_callout(
        doc,
        [content["preread"][f"status_{language}"], "0/116 | EXP-005 0/24 | מוכנות רפואית 0/6" if rtl else "0/116 | EXP-005 0/24 | Medical readiness 0/6"],
        rtl=rtl,
    )
    add_heading(doc, "מטרת הפגישה" if rtl else "Meeting purpose", level=1, rtl=rtl)
    add_body(doc, content["preread"][f"purpose_{language}"], rtl=rtl)
    add_heading(doc, "מה הוכן" if rtl else "What is prepared", level=1, rtl=rtl)
    for item in content["preread"][f"completed_{language}"]:
        add_bullet(doc, item, rtl=rtl)
    add_heading(doc, "מהלך מוצע" if rtl else "Proposed walkthrough", level=1, rtl=rtl)
    for index, item in enumerate(content["preread"][f"walkthrough_{language}"], start=1):
        add_bullet(doc, item, rtl=rtl, number=index)
    decisions_heading = add_heading(
        doc, "החלטות נדרשות" if rtl else "Decisions requested", level=1, rtl=rtl
    )
    if rtl:
        decisions_heading.paragraph_format.page_break_before = True
    for row in content["decisions"][:4]:
        question = row["question_he"] if rtl else row["question_en"]
        add_bullet(doc, f"{row['canonical_control_id']} | {question}", rtl=rtl)
    add_heading(doc, "משימת המשך מומלצת" if rtl else "Recommended next task", level=1, rtl=rtl)
    add_body(doc, content["preread"][f"recommended_next_task_{language}"], rtl=rtl)
    p = doc.add_paragraph()
    add_text_runs(p, "מקורות: " if rtl else "Sources: ", rtl=rtl)
    add_hyperlink(p, "ACL 2026 survey", SURVEY_URL)
    add_text_runs(p, " | ", rtl=rtl)
    add_hyperlink(p, "Pinned taxonomy repository", REPOSITORY_URL)
    add_text_runs(p, " | ", rtl=rtl)
    add_workbook_reference(p, content, rtl=rtl)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_tracker(content: dict[str, Any], output: Path) -> None:
    doc = Document()
    configure_document(doc, rtl=False, running_label="Supervisor tracker and decisions | 19 August 2026")
    add_title_block(
        doc,
        "Supervisor tracker and decisions",
        "Ready for Ali review - not delivered | מוכן לבדיקת עלי - לא נמסר",
        rtl=False,
        secondary_title="מעקב והחלטות למנחים",
    )
    add_evidence_callout(
        doc,
        [
            "Release gate | שער שחרור: BLOCKED pending Ali exact-package approval.",
            "Drive state | מצב Drive: link-public root; Iris Writer; Arnon absent; no permission change or access test performed.",
            "Evidence | ראיות: human screening 0/116; EXP-005 0/24; medical readiness 0/6.",
        ],
        rtl=False,
    )
    add_heading(doc, "Control status | מצב בקרות", level=1, rtl=False)
    tracker_rows = [
        [row["control"], row["status"], f"EN: {row['evidence_en']}\nHE: {row['evidence_he']}"]
        for row in content["tracker"]
    ]
    add_table(doc, ["Control | בקרה", "Status | מצב", "Evidence | ראיה"], tracker_rows, rtl=False)
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Open decisions | החלטות פתוחות", level=1, rtl=False)
    decision_rows = [
        [
            row["canonical_control_id"],
            row["state"],
            f"EN: {row['question_en']}\nHE: {row['question_he']}",
        ]
        for row in content["decisions"]
    ]
    add_table(doc, ["Canonical control ID", "State | מצב", "Question | שאלה"], decision_rows, rtl=False)
    recording_section = doc.add_section(WD_SECTION.NEW_PAGE)
    recording_section.orientation = WD_ORIENT.LANDSCAPE
    recording_section.page_width = Inches(11)
    recording_section.page_height = Inches(8.5)
    recording_section.top_margin = Inches(0.6)
    recording_section.right_margin = Inches(0.5)
    recording_section.bottom_margin = Inches(0.6)
    recording_section.left_margin = Inches(0.5)
    add_heading(
        doc,
        "Decision recording worksheet | גיליון רישום החלטות",
        level=1,
        rtl=False,
    )
    add_body(
        doc,
        "Complete every row during the meeting. Record the selected outcome, exact correction/read-back, approver and Asia/Jerusalem timestamp, owner and due date, and evidence link. All fields currently remain Pending; silence is Defer. | יש להשלים כל שורה במהלך הפגישה ולתעד תוצאה, תיקון והקראה חוזרת, מאשר/ת וחותמת זמן, אחראי/ת ומועד יעד וקישור לראיה. כל השדות ממתינים; שתיקה פירושה דחייה.",
        rtl=False,
    )
    pending = "Pending | ממתין"
    recording_rows = [
        [row["canonical_control_id"], pending, pending, pending, pending, pending]
        for row in content["decisions"]
    ]
    add_table(
        doc,
        [
            "Canonical control ID",
            "Selected outcome | תוצאה שנבחרה",
            "Correction / read-back | תיקון / הקראה חוזרת",
            "Approver / timestamp | מאשר/ת / חותמת זמן",
            "Owner / due date | אחראי/ת / מועד יעד",
            "Evidence link | קישור לראיה",
        ],
        recording_rows,
        rtl=False,
        widths_dxa=[1800, 2300, 3000, 2500, 2500, 2300],
        font_size=7.5,
    )
    add_heading(doc, "Allowed meeting outcomes | תוצאות מותרות", level=2, rtl=False)
    add_body(
        doc,
        "Confirm | Confirm with correction | Retire or supersede | Defer. Silence is Defer; no row is accepted by implication.",
        rtl=False,
    )
    p = doc.add_paragraph()
    add_text_runs(p, "Source links | קישורי מקור: ", rtl=False)
    add_hyperlink(p, "ACL survey", SURVEY_URL)
    add_text_runs(p, " | ", rtl=False)
    add_hyperlink(p, "Pinned repository", REPOSITORY_URL)
    add_text_runs(p, " | ", rtl=False)
    add_workbook_reference(p, content, rtl=False, bilingual=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def normalize_office_zip(path: Path) -> None:
    temp = path.with_suffix(path.suffix + ".normalized")
    fixed_time = (2026, 8, 19, 9, 0, 0)
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temp, "w", zipfile.ZIP_DEFLATED) as target:
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, fixed_time)
            original = source.getinfo(name)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = original.external_attr
            info.create_system = original.create_system
            target.writestr(info, source.read(name))
    os.replace(temp, path)


def ensure_node_modules(node_modules: Path, powershell_exe: Path) -> None:
    local = ROOT / "node_modules"
    if local.exists():
        return
    command = [
        str(powershell_exe),
        "-NoProfile",
        "-NonInteractive",
        "-Command",
        f"New-Item -ItemType Junction -Path '{local}' -Target '{node_modules}' | Out-Null",
    ]
    run_checked(command)


def pdf_pages(path: Path) -> int:
    return len(PdfReader(str(path)).pages)


def create_slide_pdf(preview_png: Path, output_pdf: Path) -> None:
    """Create a 16:9 PDF from the artifact-tool render without Office COM."""
    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    temporary_pdf = output_pdf.with_suffix(".tmp.pdf")
    with Image.open(preview_png) as source:
        rgb = source.convert("RGB")
        rgb.save(
            temporary_pdf,
            format="PDF",
            resolution=96.0,
            title="Human-agent classification - bilingual",
            author="Ali Hamed",
        )
    os.replace(temporary_pdf, output_pdf)


def matching_rendered_pages(render_dir: Path, pdf_stem: str) -> list[Path]:
    page_pattern = re.compile(rf"{re.escape(pdf_stem)}-(\d+)\.png")
    return sorted(
        (path for path in render_dir.iterdir() if page_pattern.fullmatch(path.name)),
        key=lambda path: int(page_pattern.fullmatch(path.name).group(1)),
    )


def render_pdf(path: Path, override_bin: Path) -> list[Path]:
    prefix = RENDER_DIR / path.stem
    executable = resolve_pdftoppm(override_bin)
    run_checked([str(executable), "-png", "-r", "144", str(path), str(prefix)])
    return matching_rendered_pages(RENDER_DIR, path.stem)


def verify_acl_source_commit(commit: str, source_manifest: Path) -> None:
    if not re.fullmatch(r"[0-9a-f]{40}", commit):
        raise RuntimeError("ACL corpus commit must be a full lowercase 40-character SHA")
    relative = source_manifest.relative_to(ROOT).as_posix()
    result = subprocess.run(
        ["git", "show", f"{commit}:{relative}"],
        cwd=ROOT,
        capture_output=True,
    )
    if result.returncode:
        raise RuntimeError("ACL corpus commit does not bind the source manifest")
    if hashlib.sha256(result.stdout).hexdigest() != sha256_file(source_manifest):
        raise RuntimeError("current ACL source manifest differs from the supplied freeze commit")


def build_phase(
    runtime: BuildRuntime,
    acl_corpus_commit: str,
    *,
    private_workbook_url: str | None,
    private_binding_sha256: str | None,
) -> None:
    content = load_content()
    import pypdf
    from validate_aug19_supervisor_package import (
        validate_content_contract,
        validate_pypdf_runtime,
    )

    validate_content_contract(content)
    validate_pypdf_runtime(pypdf.__version__, pypdf.__version__)
    if private_workbook_url:
        content["_private_workbook_url"] = private_workbook_url
    source_manifest = ROOT / "literature/acl2026-human-agent-corpus/source-manifest.json"
    verify_acl_source_commit(acl_corpus_commit, source_manifest)
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    QA_DIR.mkdir(parents=True, exist_ok=True)
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    if RENDER_DIR.exists():
        shutil.rmtree(RENDER_DIR)
    RENDER_DIR.mkdir(parents=True)

    build_chapter2(content, "en", FINAL_DIR / REQUIRED_FINAL_FILES[0])
    build_chapter2(content, "he", FINAL_DIR / REQUIRED_FINAL_FILES[2])
    build_preread(content, "en", WORK_DIR / "Supervisor_PreRead_EN.docx")
    build_preread(content, "he", WORK_DIR / "Supervisor_PreRead_HE.docx")
    build_tracker(content, WORK_DIR / "Supervisor_Tracker_and_Decisions_Bilingual.docx")

    for docx_path in (
        FINAL_DIR / "Chapter_2_Literature_Review_EN.docx",
        FINAL_DIR / "Chapter_2_Literature_Review_HE.docx",
        WORK_DIR / "Supervisor_PreRead_EN.docx",
        WORK_DIR / "Supervisor_PreRead_HE.docx",
        WORK_DIR / "Supervisor_Tracker_and_Decisions_Bilingual.docx",
    ):
        run_checked([str(runtime.python_exe), str(runtime.docx_sanitizer), str(docx_path), "--in-place"])
        normalize_office_zip(docx_path)

    ensure_node_modules(runtime.node_modules, runtime.powershell_exe)
    pptx_work = WORK_DIR / "Human_Agent_Classification_Bilingual.pptx"
    run_checked(
        [
            str(runtime.node_exe),
            str(ROOT / "scripts/build_aug19_classification_slide.mjs"),
            "--content",
            str(CONTENT_PATH),
            "--output",
            str(pptx_work),
            "--preview",
            str(RENDER_DIR / "Human_Agent_Classification_Bilingual-artifact-tool.png"),
            "--layout",
            str(WORK_DIR / "Human_Agent_Classification_Bilingual.layout.json"),
        ]
    )
    normalize_office_zip(pptx_work)
    shutil.copyfile(pptx_work, FINAL_DIR / "Human_Agent_Classification_Bilingual.pptx")
    create_slide_pdf(
        RENDER_DIR / "Human_Agent_Classification_Bilingual-artifact-tool.png",
        FINAL_DIR / "Human_Agent_Classification_Bilingual.pdf",
    )

    office_receipt = WORK_DIR / "office-export-receipt.json"
    run_checked(
        [
            str(runtime.powershell_exe),
            "-NoProfile",
            "-NonInteractive",
            "-ExecutionPolicy",
            "Bypass",
            "-File",
            str(ROOT / "scripts/export_aug19_office_artifacts.ps1"),
            "-PackageDirectory",
            str(FINAL_DIR),
            "-BuildWorkDirectory",
            str(WORK_DIR),
            "-ReceiptPath",
            str(office_receipt),
            "-TimeoutSeconds",
            "120",
            "-MaxAttempts",
            "2",
        ]
    )

    render_items: list[dict[str, Any]] = []
    for pdf_name in (
        "Chapter_2_Literature_Review_EN.pdf",
        "Chapter_2_Literature_Review_HE.pdf",
        "Human_Agent_Classification_Bilingual.pdf",
        "Supervisor_PreRead_EN.pdf",
        "Supervisor_PreRead_HE.pdf",
        "Supervisor_Tracker_and_Decisions_Bilingual.pdf",
    ):
        pdf_path = FINAL_DIR / pdf_name
        images = render_pdf(pdf_path, runtime.override_bin)
        if len(images) != pdf_pages(pdf_path):
            raise RuntimeError(f"render count mismatch for {pdf_name}")
        render_items.extend(
            {
                "artifact": pdf_name,
                "page_or_slide": index,
                "render_file": image.name,
                "sha256": sha256_file(image),
            }
            for index, image in enumerate(images, start=1)
        )
    artifact_preview = RENDER_DIR / "Human_Agent_Classification_Bilingual-artifact-tool.png"
    render_items.append(
        {
            "artifact": "Human_Agent_Classification_Bilingual.pptx",
            "page_or_slide": 1,
            "render_file": artifact_preview.name,
            "sha256": sha256_file(artifact_preview),
        }
    )
    write_text_lf(
        WORK_DIR / "render-index.json",
        json.dumps({"schema_version": "vego-ai.render-index.v1", "items": render_items}, indent=2),
    )

    import docx

    artifact_package = json.loads(
        (runtime.node_modules / "@oai/artifact-tool/package.json").read_text(encoding="utf-8")
    )
    node_version = subprocess.check_output([str(runtime.node_exe), "--version"], text=True).strip()
    powershell_version = subprocess.check_output(
        [
            str(runtime.powershell_exe),
            "-NoProfile",
            "-NonInteractive",
            "-Command",
            "$PSVersionTable.PSVersion.ToString()",
        ],
        text=True,
    ).strip()
    office_data = json.loads(office_receipt.read_text(encoding="utf-8-sig"))
    control_register = ROOT / "docs/research/meetings/2026-08-12-control-register.csv"
    build_receipt = {
        "schema_version": "vego-ai.aug19-package-build-receipt.v1",
        "package_status": "Ready for Ali review - not delivered, accepted, or closed",
        "sources": {
            "content_sha256": sha256_file(CONTENT_PATH),
            "control_register_sha256": sha256_file(control_register),
            "acl_source_manifest_sha256": sha256_file(source_manifest),
            "acl_corpus_git_commit": acl_corpus_commit,
            "taxonomy_repository_commit": TAXONOMY_REPOSITORY_COMMIT,
        },
        "runtimes": {
            "python": sys.version.split()[0],
            "python_executable_sha256": sha256_file(runtime.python_exe),
            "python_docx": docx.__version__,
            "pypdf": pypdf.__version__,
            "node": node_version,
            "node_executable_sha256": sha256_file(runtime.node_exe),
            "powershell": powershell_version,
            "powershell_executable_sha256": sha256_file(runtime.powershell_exe),
            "artifact_tool": artifact_package["version"],
            "office": sorted({row["office_version"] for row in office_data["exports"]}),
            "pdf_renderer": "bundled pdftoppm",
            "pdf_renderer_sha256": sha256_file(resolve_pdftoppm(runtime.override_bin)),
            "docx_sanitizer_sha256": sha256_file(runtime.docx_sanitizer),
        },
        "office_export_receipt_sha256": sha256_file(office_receipt),
        "artifacts_before_manifest": {
            name: {
                "sha256": sha256_file(FINAL_DIR / name),
                "bytes": (FINAL_DIR / name).stat().st_size,
                "pages": pdf_pages(FINAL_DIR / name) if name.endswith(".pdf") else None,
            }
            for name in REQUIRED_FINAL_FILES
            if name != "package-manifest.sanitized.json" and (FINAL_DIR / name).exists()
        },
        "render_index_sha256": sha256_file(WORK_DIR / "render-index.json"),
        "visual_inspection": "Pending",
        "workbook_binding": {
            "logical_alias": content["workbook"]["logical_alias"],
            "mode": "ali-approved-private-binding" if private_workbook_url else "withheld",
            "binding_sha256": private_binding_sha256,
            "live_url_recorded": False,
        },
    }
    write_text_lf(
        QA_DIR / "build-receipt.json",
        json.dumps(build_receipt, indent=2, ensure_ascii=False) + "\n",
    )


def manifest_phase() -> None:
    content = load_content()
    build_receipt_path = QA_DIR / "build-receipt.json"
    render_receipt_path = QA_DIR / "render-receipt.json"
    if not build_receipt_path.is_file() or not render_receipt_path.is_file():
        raise RuntimeError("build and render receipts must exist before manifest freeze")
    build_receipt = json.loads(build_receipt_path.read_text(encoding="utf-8"))
    render_receipt = json.loads(render_receipt_path.read_text(encoding="utf-8"))
    import pypdf
    from validate_aug19_supervisor_package import validate_pypdf_runtime

    validate_pypdf_runtime(
        pypdf.__version__, str(build_receipt.get("runtimes", {}).get("pypdf", ""))
    )
    if render_receipt.get("all_pages_and_slides_visually_inspected") is not True:
        raise RuntimeError("visual inspection receipt is not complete")
    artifacts = []
    for name in REQUIRED_FINAL_FILES:
        if name == "package-manifest.sanitized.json":
            continue
        path = FINAL_DIR / name
        if not path.is_file():
            raise RuntimeError(f"missing final artifact: {name}")
        row = {
            "file": name,
            "sha256": sha256_file(path),
            "bytes": path.stat().st_size,
            "media_type": {
                ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
                ".pdf": "application/pdf",
            }[path.suffix.lower()],
        }
        if path.suffix.lower() == ".pdf":
            row["pages"] = pdf_pages(path)
        elif path.suffix.lower() == ".pptx":
            row["slides"] = 1
        artifacts.append(row)

    sources = build_receipt["sources"]
    runtimes = build_receipt["runtimes"]
    manifest = {
        "schema_version": "vego-ai.aug19-supervisor-package-manifest.v1",
        "package_status": "Ready for Ali review - not delivered, accepted, or closed",
        "meeting": content["meeting"],
        "scope": "Literature only",
        "artifacts": artifacts,
        "bindings": {
            "acl_corpus_git_commit": sources["acl_corpus_git_commit"],
            "taxonomy_repository_commit": sources["taxonomy_repository_commit"],
            "content_sha256": sources["content_sha256"],
            "control_register_sha256": sources["control_register_sha256"],
            "acl_source_manifest_sha256": sources["acl_source_manifest_sha256"],
            "python_runtime": runtimes["python"],
            "python_executable_sha256": runtimes["python_executable_sha256"],
            "python_docx_runtime": runtimes["python_docx"],
            "pypdf_runtime": runtimes["pypdf"],
            "node_runtime": runtimes["node"],
            "node_executable_sha256": runtimes["node_executable_sha256"],
            "powershell_runtime": runtimes["powershell"],
            "powershell_executable_sha256": runtimes["powershell_executable_sha256"],
            "artifact_tool_runtime": runtimes["artifact_tool"],
            "office_runtime": ",".join(runtimes["office"]),
            "pdf_renderer_sha256": runtimes["pdf_renderer_sha256"],
            "docx_sanitizer_sha256": runtimes["docx_sanitizer_sha256"],
            "build_receipt_sha256": sha256_file(build_receipt_path),
            "render_receipt_sha256": sha256_file(render_receipt_path),
        },
        "boundaries": content["boundaries"],
        "workbook": {
            "title": content["workbook"]["title"],
            "logical_alias": content["workbook"]["logical_alias"],
            "delivery": content["workbook"]["delivery"],
            "link_status": content["workbook"]["link_status"],
            "binding_mode": build_receipt["workbook_binding"]["mode"],
            "release_status": content["workbook"]["release_status"],
        },
        "release": {
            "ali_exact_package_approval": "Pending",
            "drive_permission_correction": "Blocked pending Ali approval",
            "recipient_access_tests": "Not performed",
            "delivered": False,
        },
        "qa": {
            "docx_pdf_page_parity": render_receipt["docx_pdf_page_parity"],
            "all_pages_and_slides_visually_inspected": True,
            "english_hebrew_parity": render_receipt["english_hebrew_parity"],
            "rtl_inspected": render_receipt["rtl_inspected"],
            "citations_and_links_inspected": render_receipt["citations_and_links_inspected"],
        },
    }
    write_text_lf(
        FINAL_DIR / "package-manifest.sanitized.json",
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
    )


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--phase", choices=("build", "manifest"), default="build")
    parser.add_argument("--python-exe", type=Path)
    parser.add_argument("--powershell-exe", type=Path)
    parser.add_argument("--node-exe", type=Path)
    parser.add_argument("--node-modules", type=Path)
    parser.add_argument("--override-bin", type=Path)
    parser.add_argument("--docx-sanitizer", type=Path)
    parser.add_argument("--acl-corpus-commit")
    parser.add_argument("--private-workbook-binding", type=Path)
    parser.add_argument("--ali-approved-release", action="store_true")
    args = parser.parse_args()
    if args.phase == "build":
        missing = [
            name
            for name in (
                "python_exe",
                "powershell_exe",
                "node_exe",
                "node_modules",
                "override_bin",
                "docx_sanitizer",
                "acl_corpus_commit",
            )
            if getattr(args, name) is None
        ]
        if missing:
            parser.error(f"build phase requires: {', '.join('--' + name.replace('_', '-') for name in missing)}")
        runtime = BuildRuntime(
            python_exe=require_runtime_path(args.python_exe),
            powershell_exe=require_runtime_path(args.powershell_exe),
            node_exe=require_runtime_path(args.node_exe),
            node_modules=require_runtime_path(args.node_modules, directory=True),
            override_bin=require_runtime_path(args.override_bin, directory=True),
            docx_sanitizer=require_runtime_path(args.docx_sanitizer),
        )
        resolve_pdftoppm(runtime.override_bin)
        require_runtime_path(runtime.node_modules / "@oai/artifact-tool/package.json")
        content = load_content()
        private_workbook_url, private_binding_sha256 = load_private_workbook_binding(
            content,
            args.private_workbook_binding,
            ali_approved_release=args.ali_approved_release,
        )
        build_phase(
            runtime,
            args.acl_corpus_commit,
            private_workbook_url=private_workbook_url,
            private_binding_sha256=private_binding_sha256,
        )
    else:
        manifest_phase()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
