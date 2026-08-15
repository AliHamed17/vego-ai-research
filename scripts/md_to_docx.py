#!/usr/bin/env python3
"""Convert the proposal Markdown files to supervisor-ready Word documents.

Deliberately small and predictable: handles the subset of Markdown these
documents actually use - ATX headings, paragraphs with inline bold/italic/code,
blockquotes, bullet and numbered lists, pipe tables, and horizontal rules.

Usage:  python scripts/md_to_docx.py <input.md> <output.docx> [--title "Doc title"]
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ACCENT = RGBColor(0x1F, 0x38, 0x64)
INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", re.DOTALL)


def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def add_inline(par, text: str, base_bold: bool = False, base_italic: bool = False) -> None:
    """Render inline **bold**, *italic* and `code` inside one paragraph."""
    text = re.sub(r"\[([^\]]+)\]\((?:[^)]+)\)", r"\1", text)  # links -> label
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = par.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = par.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x66, 0x00, 0x99)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = par.add_run(part[1:-1])
            run.italic = True
        else:
            run = par.add_run(part)
        if base_bold:
            run.bold = True
        if base_italic:
            run.italic = True


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:|-]+\|?", line.strip())) and "-" in line


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(ncols):
            text = row[j] if j < len(row) else ""
            par = cells[j].paragraphs[0]
            add_inline(par, text, base_bold=(i == 0))
            for run in par.runs:
                run.font.size = Pt(8.5)
                if i == 0:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if i == 0:
                shade(cells[j], "1F3864")
            elif i % 2 == 0:
                shade(cells[j], "F2F5FA")
    doc.add_paragraph()


def add_quote(doc: Document, lines: list[str]) -> None:
    for line in lines:
        par = doc.add_paragraph()
        par.paragraph_format.left_indent = Pt(24)
        par.paragraph_format.space_after = Pt(3)
        add_inline(par, line)
        for run in par.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()


def convert(md_path: Path, docx_path: Path, title: str | None) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)

    for level, size in ((1, 19), (2, 14.5), (3, 12)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = ACCENT
        st.font.bold = True

    if title:
        head = doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = head.add_run(title)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = ACCENT
        doc.add_paragraph()

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i, n = 0, len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if set(stripped) <= {"-", "*", "_"} and len(stripped) >= 3:
            par = doc.add_paragraph()
            par.paragraph_format.space_before = Pt(4)
            bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "BFBFBF")
            bdr.append(bottom)
            par._p.get_or_add_pPr().append(bdr)
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            text = re.sub(r"[*`]", "", m.group(2)).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        if stripped.startswith(">"):
            block = []
            while i < n and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            add_quote(doc, [b for b in block if b])
            continue

        if stripped.startswith("|"):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cur = lines[i].strip()
                if not is_separator(cur):
                    rows.append(split_row(cur))
                i += 1
            add_table(doc, rows)
            continue

        m = re.match(r"^(\s*)([-*+])\s+(.*)$", raw)
        if m:
            indent = len(m.group(1))
            style = "List Bullet" if indent < 2 else "List Bullet 2"
            par = doc.add_paragraph(style=style)
            add_inline(par, m.group(3))
            i += 1
            continue

        m = re.match(r"^(\s*)(\d+)[.)]\s+(.*)$", raw)
        if m:
            par = doc.add_paragraph(style="List Number")
            add_inline(par, m.group(3))
            i += 1
            continue

        # paragraph: join soft-wrapped lines
        block = []
        while i < n:
            cur = lines[i].rstrip()
            s = cur.strip()
            if (not s or s.startswith(("#", ">", "|", "---"))
                    or re.match(r"^\s*([-*+]|\d+[.)])\s+", cur)):
                break
            block.append(s)
            i += 1
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(par, " ".join(block))

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    ap.add_argument("--title", default=None)
    args = ap.parse_args()
    convert(args.input, args.output, args.title)
    print(f"docx -> {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
