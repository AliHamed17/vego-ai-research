"""Build the 3 September 2026 VEGO-AI supervisor DOCX package.

The source is the sanitized Markdown package. Figures are generated from aggregate
values embedded below; no private student artifact or reviewer text is read.
"""

from __future__ import annotations

import argparse
import os
import re
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFilter, ImageFont

ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "docs/research/phd-proposal/2026-09-03-supervisor-review-package"
DEFAULT_OUTPUT = ROOT / "output/doctoral-proposal/2026-09-03"

NAVY = "102A43"
BLUE = "1F6AA5"
CYAN = "31A6B8"
TEAL = "2A9D8F"
GOLD = "E9B949"
RED = "C14953"
INK = "1E2933"
MUTED = "5B6773"
PALE = "F3F7FA"
WHITE = "FFFFFF"
LIGHT_BLUE = "EAF3F9"
LIGHT_GOLD = "FFF6DD"
LIGHT_RED = "FCECEE"


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_directory = Path(os.environ.get("WINDIR", "")) / "Fonts"
    candidates = [
        font_directory / "aptos.ttf",
        font_directory / "arial.ttf",
        font_directory / "segoeui.ttf",
        font_directory / "calibri.ttf",
    ]
    if bold:
        candidates = [
            font_directory / "aptos-bold.ttf",
            font_directory / "arialbd.ttf",
            font_directory / "seguisb.ttf",
            font_directory / "calibrib.ttf",
        ] + candidates
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _gradient(width: int, height: int) -> Image.Image:
    top = (248, 251, 253)
    bottom = (229, 241, 247)
    image = Image.new("RGB", (width, height), top)
    draw = ImageDraw.Draw(image)
    for y in range(height):
        t = y / max(1, height - 1)
        color = tuple(round(a + (b - a) * t) for a, b in zip(top, bottom, strict=True))
        draw.line((0, y, width, y), fill=color)
    return image


def _wrap(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _card(
    canvas: Image.Image,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    accent: str,
    *,
    title_size: int = 38,
    body_size: int = 29,
) -> None:
    x0, y0, x1, y1 = box
    shadow = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    sd = ImageDraw.Draw(shadow)
    sd.rounded_rectangle((x0 + 12, y0 + 16, x1 + 12, y1 + 16), 32, fill=(9, 34, 54, 42))
    canvas.alpha_composite(shadow.filter(ImageFilter.GaussianBlur(12)))
    draw = ImageDraw.Draw(canvas)
    draw.rounded_rectangle(box, 32, fill=(255, 255, 255, 245), outline="#C9D8E2", width=3)
    draw.rounded_rectangle((x0, y0, x0 + 18, y1), 22, fill=f"#{accent}")
    title_font = _font(title_size, bold=True)
    body_font = _font(body_size)
    draw.text((x0 + 48, y0 + 34), title, font=title_font, fill=f"#{NAVY}")
    y = y0 + 96
    for line in _wrap(draw, body, body_font, x1 - x0 - 88):
        draw.text((x0 + 48, y), line, font=body_font, fill=f"#{INK}")
        y += body_size + 13


def _title(draw: ImageDraw.ImageDraw, title: str, subtitle: str) -> None:
    draw.text((120, 74), title, font=_font(58, True), fill=f"#{NAVY}")
    draw.text((122, 150), subtitle, font=_font(30), fill=f"#{MUTED}")
    draw.rounded_rectangle((120, 211, 2280, 219), 4, fill=f"#{CYAN}")


def _arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line((*start, *end), fill=f"#{BLUE}", width=8)
    x, y = end
    draw.polygon([(x, y), (x - 25, y - 16), (x - 25, y + 16)], fill=f"#{BLUE}")


def build_figures(directory: Path) -> dict[str, Path]:
    directory.mkdir(parents=True, exist_ok=True)
    figures: dict[str, Path] = {}

    # Figure 1: VEGO-AI architecture and human decision points.
    image = _gradient(2400, 1350).convert("RGBA")
    draw = ImageDraw.Draw(image)
    _title(
        draw,
        "VEGO-AI decision path",
        "Four agent stages; a human decision object can be created at each boundary",
    )
    cards = [
        ("1  LANGUAGE", "Template and notation questions", BLUE),
        ("2  DOMAIN", "Reference guideline construction", CYAN),
        ("3  INSPECT", "Claim-level model assessment", TEAL),
        ("4  VARIABILITY", "Cross-case pattern classification", GOLD),
    ]
    x_positions = [120, 690, 1260, 1830]
    for index, ((heading, body, accent), x) in enumerate(zip(cards, x_positions, strict=True)):
        _card(image, (x, 330, x + 450, 640), heading, body, accent, title_size=31, body_size=27)
        if index < 3:
            _arrow(draw, (x + 462, 485), (x_positions[index + 1] - 18, 485))
    draw.rounded_rectangle((120, 760, 2280, 1190), 34, fill="#102A43", outline="#1F6AA5", width=4)
    draw.text((180, 808), "GOVERNED HUMAN LAYER", font=_font(42, True), fill="#FFFFFF")
    gate_font = _font(28, True)
    small = _font(25)
    gates = [
        ("DETECT", "uncertainty · omission · conflict"),
        ("ROUTE", "competence ≠ authority"),
        ("RECORD", "reason · evidence · dissent"),
        ("REUSE", "authorize · validate · receipt"),
    ]
    for i, (head, body) in enumerate(gates):
        x = 180 + i * 520
        draw.rounded_rectangle(
            (x, 910, x + 430, 1110), 26, fill="#1A3D5A", outline="#5BC0CE", width=3
        )
        draw.text((x + 30, 946), head, font=gate_font, fill="#69D2DC")
        for n, line in enumerate(_wrap(draw, body, small, 360)):
            draw.text((x + 30, 1004 + n * 36), line, font=small, fill="#FFFFFF")
    path = directory / "fig-01-vego-human-touchpoints.png"
    image.convert("RGB").save(path, dpi=(300, 300), quality=96)
    figures["architecture"] = path

    # Figure 2: residual gap.
    image = _gradient(2400, 1350).convert("RGBA")
    draw = ImageDraw.Draw(image)
    _title(
        draw,
        "Residual research gap",
        "Established components are conceded; the claim is an evaluated integration",
    )
    blocks = [
        ("ROUTE", "Selective prediction\nLearning to defer\nExpert matching", BLUE),
        ("RECORD", "Provenance\nReviewability\nTruth maintenance", CYAN),
        ("REUSE", "Corrective memory\nPrecedent\nAuthorization + transfer", TEAL),
    ]
    for i, (head, body, accent) in enumerate(blocks):
        x = 120 + i * 760
        _card(
            image,
            (x, 320, x + 640, 790),
            head,
            body.replace("\n", " · "),
            accent,
            title_size=40,
            body_size=31,
        )
    draw.rounded_rectangle((250, 900, 2150, 1195), 44, fill="#102A43", outline="#E9B949", width=5)
    draw.text((335, 946), "THE TESTABLE THREE-WAY CLAIM", font=_font(42, True), fill="#F4CB66")
    claim = "claim-level routing by competence + authority  •  reconstructable lifecycle record  •  authorized, scope-aware reuse"
    y = 1022
    for line in _wrap(draw, claim, _font(30), 1720):
        draw.text((335, y), line, font=_font(30), fill="#FFFFFF")
        y += 48
    path = directory / "fig-02-residual-gap.png"
    image.convert("RGB").save(path, dpi=(300, 300), quality=96)
    figures["gap"] = path

    # Figure 3: prospective Study 1 flow.
    image = _gradient(2400, 1350).convert("RGBA")
    draw = ImageDraw.Draw(image)
    _title(
        draw,
        "Study 1 controlled benchmark",
        "Same events, same evidence, same budget; competence and authority remain separate",
    )
    flow = [
        ("FREEZE", "events · splits · labels"),
        ("QUALIFY", "claim-type competence"),
        ("AUTHORIZE", "role / mandate matrix"),
        ("REPLAY", "7 matched-budget arms"),
        ("EVALUATE", "capture + conditional correctness"),
    ]
    widths = [370, 370, 370, 370, 430]
    xs = [90, 545, 1000, 1455, 1910]
    accents = [BLUE, CYAN, GOLD, TEAL, RED]
    for i, ((head, body), x, width, accent) in enumerate(
        zip(flow, xs, widths, accents, strict=True)
    ):
        _card(image, (x, 360, x + width, 690), head, body, accent, title_size=33, body_size=25)
        if i < len(flow) - 1:
            _arrow(draw, (x + width + 8, 525), (xs[i + 1] - 18, 525))
    draw.rounded_rectangle((150, 840, 2250, 1175), 34, fill="#FFFFFF", outline="#B9CBD6", width=3)
    draw.text((210, 882), "HARD GATES", font=_font(38, True), fill=f"#{RED}")
    gates = [
        "missing authority → block",
        "revoked evidence → zero influence",
        "budget exceeded → defer",
        "independent labels missing → no outcome claim",
    ]
    for i, item in enumerate(gates):
        y = 958 + (i // 2) * 86
        x = 210 + (i % 2) * 1010
        draw.ellipse((x, y + 6, x + 28, y + 34), fill=f"#{RED}")
        draw.text((x + 48, y), item, font=_font(28), fill=f"#{INK}")
    path = directory / "fig-03-study1-benchmark.png"
    image.convert("RGB").save(path, dpi=(300, 300), quality=96)
    figures["study1"] = path

    # Figure 4: preliminary evidence dashboard.
    image = _gradient(2400, 1350).convert("RGBA")
    draw = ImageDraw.Draw(image)
    _title(
        draw,
        "Preliminary evidence",
        "Descriptive and technical feasibility only — no human-benefit or accuracy claim",
    )
    panels = [
        ("STAGE SIGNALS", "6 / 18 / 506 / 11", "candidate signals by stages 1–4", BLUE),
        ("RECORDED REVIEW", "108 of 120", "changes included by non-Satisfied rule", CYAN),
        ("C0 REPLAY", "0 of 1,874", "joint-policy selections; required signals absent", GOLD),
        ("ONE CORRECTION", "17.5 → 16.5", "deterministic score propagation", TEAL),
    ]
    for i, (head, number, body, accent) in enumerate(panels):
        x = 110 + (i % 2) * 1150
        y = 310 + (i // 2) * 480
        _card(image, (x, y, x + 1030, y + 390), head, body, accent, title_size=34, body_size=28)
        draw.text((x + 500, y + 125), number, font=_font(58, True), fill=f"#{accent}")
    draw.text(
        (120, 1254),
        "Interpretation depends on denominators, selection process, and evidence state shown in the text.",
        font=_font(25),
        fill=f"#{MUTED}",
    )
    path = directory / "fig-04-preliminary-evidence.png"
    image.convert("RGB").save(path, dpi=(300, 300), quality=96)
    figures["results"] = path

    # Figure 5: programme dependency.
    image = _gradient(2400, 1350).convert("RGBA")
    draw = ImageDraw.Draw(image)
    _title(
        draw,
        "Research programme",
        "Each study can fail; the integrated study answers the umbrella question",
    )
    entries = [
        ("STUDY 1", "When + whom", "routing policy", BLUE),
        ("STUDY 2", "What record", "judgment contract", CYAN),
        ("STUDY 3", "When reusable", "reuse gate", TEAL),
    ]
    for i, (head, sub, body, accent) in enumerate(entries):
        x = 150 + i * 760
        _card(
            image,
            (x, 320, x + 610, 730),
            head,
            f"{sub}. Artifact: {body}.",
            accent,
            title_size=41,
            body_size=31,
        )
        _arrow(draw, (x + 305, 748), (1200, 900))
    draw.rounded_rectangle((510, 900, 1890, 1190), 44, fill="#102A43", outline="#E9B949", width=5)
    draw.text((665, 952), "INTEGRATED FOUR-ARM TEST", font=_font(45, True), fill="#F4CB66")
    for n, line in enumerate(
        _wrap(
            draw,
            "Gated objective: authority + revocation + scope safety, then correctness at fixed attention",
            _font(29),
            1190,
        )
    ):
        draw.text((640, 1030 + n * 48), line, font=_font(29), fill="#FFFFFF")
    path = directory / "fig-05-programme.png"
    image.convert("RGB").save(path, dpi=(300, 300), quality=96)
    figures["programme"] = path
    return figures


def _xml_name(prefix: str, local_name: str) -> str:
    return f"{prefix}:{local_name}"


def _qname(prefix: str, local_name: str) -> str:
    return qn(_xml_name(prefix, local_name))


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(_qname("w", "shd"))
    if shd is None:
        shd = OxmlElement(_xml_name("w", "shd"))
        tc_pr.append(shd)
    shd.set(_qname("w", "fill"), fill)


def _set_cell_margins(cell, value: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(_qname("w", "tcMar"))
    if margins is None:
        margins = OxmlElement(_xml_name("w", "tcMar"))
        tc_pr.append(margins)
    for edge in ("top", "start", "bottom", "end"):
        node = margins.find(_qname("w", edge))
        if node is None:
            node = OxmlElement(_xml_name("w", edge))
            margins.append(node)
        node.set(_qname("w", "w"), str(value))
        node.set(_qname("w", "type"), "dxa")


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(_qname("w", "tblHeader")) is None:
        tr_pr.append(OxmlElement(_xml_name("w", "tblHeader")))


def _no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(_qname("w", "cantSplit")) is None:
        tr_pr.append(OxmlElement(_xml_name("w", "cantSplit")))


def _page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld = OxmlElement(_xml_name("w", "fldSimple"))
    fld.set(_qname("w", "instr"), "PAGE")
    run._r.append(fld)


def _hyperlink(paragraph, label: str, url: str, *, size: float) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement(_xml_name("w", "hyperlink"))
    link.set(_qname("r", "id"), relationship)
    run = OxmlElement(_xml_name("w", "r"))
    properties = OxmlElement(_xml_name("w", "rPr"))
    color = OxmlElement(_xml_name("w", "color"))
    color.set(_qname("w", "val"), BLUE)
    properties.append(color)
    underline = OxmlElement(_xml_name("w", "u"))
    underline.set(_qname("w", "val"), "single")
    properties.append(underline)
    size_element = OxmlElement(_xml_name("w", "sz"))
    size_element.set(_qname("w", "val"), str(round(size * 2)))
    properties.append(size_element)
    run.append(properties)
    text = OxmlElement(_xml_name("w", "t"))
    text.text = label
    run.append(text)
    link.append(run)
    paragraph._p.append(link)


INLINE = re.compile(r"(\[[^\]]+\]\([^)]+\)|https?://[^\s<>]+|\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)")


def _inline(paragraph, text: str, *, size: float = 9.6, color: str = INK) -> None:
    text = text.replace("  ", " ").strip()
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            run.font.size = Pt(size)
            run.font.color.rgb = _rgb(color)
        token = match.group(0)
        if token.startswith("["):
            link = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            assert link
            label, target = link.groups()
            if label.startswith("`") and label.endswith("`"):
                label = label[1:-1]
            if target.startswith(("http://", "https://")):
                _hyperlink(paragraph, label, target, size=size)
            else:
                run = paragraph.add_run(label)
                run.font.size = Pt(size)
                run.font.color.rgb = _rgb(BLUE)
        elif token.startswith(("http://", "https://")):
            url = token.rstrip(".,;")
            _hyperlink(paragraph, url, url, size=size)
            suffix = token[len(url) :]
            if suffix:
                run = paragraph.add_run(suffix)
                run.font.size = Pt(size)
                run.font.color.rgb = _rgb(color)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.size = Pt(size)
            run.font.color.rgb = _rgb(color)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.size = Pt(size)
            run.font.color.rgb = _rgb(color)
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(max(7.0, size - 0.8))
            run.font.color.rgb = _rgb(NAVY)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.font.size = Pt(size)
        run.font.color.rgb = _rgb(color)


def _configure_document(doc: Document, *, compact: bool = False) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    margin = 0.47 if compact else 0.68
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.0 if compact else 10.2)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(3.4 if compact else 5.5)
    pf.line_spacing = 1.08 if compact else 1.2

    for level, size, before, after, colour in [
        (1, 16.0, 12, 6, NAVY),
        (2, 12.2, 9, 4, BLUE),
        (3, 10.4, 6, 3, NAVY),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size - (1.5 if compact else 0))
        style.font.bold = True
        style.font.color.rgb = _rgb(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" in doc.styles:
        caption = doc.styles["Caption"]
    else:
        caption = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Aptos"
    caption.font.size = Pt(7.8 if compact else 8.2)
    caption.font.italic = True
    caption.font.color.rgb = _rgb(MUTED)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("VEGO-AI  |  supervisor-review candidate  |  3 September 2026")
    run.font.name = "Aptos"
    run.font.size = Pt(7.4)
    run.font.color.rgb = _rgb(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Ready for supervisor review · not approved     ")
    run.font.size = Pt(7.2)
    run.font.color.rgb = _rgb(MUTED)
    _page_number(footer)


def _add_status_box(doc: Document, text: str, *, fill: str = LIGHT_GOLD) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.8)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, 130)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _inline(paragraph, text, size=9.0, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def _title_page(doc: Document, title: str, subtitle: str, kind: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(22)
    run = paragraph.add_run("VEGO-AI")
    run.font.name = "Aptos Display"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = _rgb(CYAN)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(36)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    run.font.name = "Aptos Display"
    run.font.size = Pt(27)
    run.font.bold = True
    run.font.color.rgb = _rgb(NAVY)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run(subtitle)
    run.font.name = "Aptos Display"
    run.font.size = Pt(14)
    run.font.color.rgb = _rgb(BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    pPr = p._p.get_or_add_pPr()
    border = OxmlElement(_xml_name("w", "pBdr"))
    bottom = OxmlElement(_xml_name("w", "bottom"))
    bottom.set(_qname("w", "val"), "single")
    bottom.set(_qname("w", "sz"), "26")
    bottom.set(_qname("w", "color"), CYAN)
    border.append(bottom)
    pPr.append(border)

    metadata = doc.add_table(rows=4, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata.autofit = False
    labels = [
        ("Document", kind),
        ("Author", "Ali Hamed"),
        ("Date", "3 September 2026"),
        ("Status", "Ready for supervisor review · decisions pending"),
    ]
    for row, (label, value) in zip(metadata.rows, labels, strict=True):
        _no_split(row)
        row.cells[0].width = Inches(1.35)
        row.cells[1].width = Inches(5.35)
        _set_cell_shading(row.cells[0], NAVY)
        _set_cell_shading(row.cells[1], PALE)
        _set_cell_margins(row.cells[0], 120)
        _set_cell_margins(row.cells[1], 120)
        p0 = row.cells[0].paragraphs[0]
        _inline(p0, label, size=8.6, color=WHITE)
        for run in p0.runs:
            run.bold = True
        _inline(row.cells[1].paragraphs[0], value, size=8.6, color=INK)

    doc.add_paragraph()
    _add_status_box(
        doc,
        "**Evidence boundary.** Preliminary results demonstrate descriptive signal availability and deterministic technical propagation. They do not demonstrate accuracy, human benefit, burden reduction, causality, policy superiority, or safe transfer.",
    )

    if kind.startswith("Doctoral"):
        table = doc.add_table(rows=2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for index, (head, body, colour) in enumerate(
            [
                ("ROUTE", "When and whom to ask", BLUE),
                ("RECORD", "What judgment must retain", CYAN),
                ("REUSE", "When an old answer may act", TEAL),
            ]
        ):
            _set_cell_shading(table.cell(0, index), colour)
            _set_cell_shading(table.cell(1, index), "F7FAFC")
            _inline(table.cell(0, index).paragraphs[0], head, size=9.0, color=WHITE)
            for run in table.cell(0, index).paragraphs[0].runs:
                run.bold = True
            _inline(table.cell(1, index).paragraphs[0], body, size=8.2, color=INK)
            _set_cell_margins(table.cell(0, index), 110)
            _set_cell_margins(table.cell(1, index), 110)
    doc.add_page_break()


def _markdown_blocks(lines: list[str]) -> Iterable[tuple[str, object]]:
    i = 0
    paragraph: list[str] = []

    def flush() -> Iterable[tuple[str, object]]:
        nonlocal paragraph
        if paragraph:
            value = " ".join(part.strip() for part in paragraph).strip()
            paragraph = []
            if value:
                yield ("paragraph", value)

    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            yield from flush()
            language = line[3:].strip()
            code: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i].rstrip("\n"))
                i += 1
            yield ("code", (language, "\n".join(code)))
        elif re.match(r"^#{1,3} ", line):
            yield from flush()
            level = len(line) - len(line.lstrip("#"))
            yield ("heading", (level, line[level + 1 :].strip()))
        elif line.startswith("|") and line.count("|") >= 2:
            yield from flush()
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                    rows.append(cells)
                i += 1
            i -= 1
            yield ("table", rows)
        elif re.match(r"^[-*] ", line):
            yield from flush()
            yield ("bullet", line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            yield from flush()
            yield ("number", re.sub(r"^\d+\. ", "", line).strip())
        elif not line.strip():
            yield from flush()
        else:
            paragraph.append(line)
        i += 1
    yield from flush()


def _add_figure(doc: Document, path: Path, caption: str, alt: str, *, width: float = 6.9) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(caption)


def _add_table(doc: Document, rows: list[list[str]], *, compact: bool = False) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.add_row()
        _no_split(row)
        if row_index == 0:
            _repeat_header(row)
        for index in range(columns):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, 62 if compact else 82)
            _set_cell_shading(
                cell, NAVY if row_index == 0 else ("F7FAFC" if row_index % 2 else WHITE)
            )
            text = values[index] if index < len(values) else ""
            _inline(
                cell.paragraphs[0],
                text,
                size=7.0 if compact else 7.25,
                color=WHITE if row_index == 0 else INK,
            )
            if row_index == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


FIGURE_MAP = {
    "1.2 VEGO-AI baseline and decision points": (
        "architecture",
        "Figure 1. VEGO-AI stages and the proposed governed human layer.",
        "Four-stage VEGO-AI architecture with detect, route, record, and reuse controls.",
    ),
    "2.7 Residual gap and contribution boundary": (
        "gap",
        "Figure 2. The residual gap after conceding established mechanisms.",
        "Three established streams feeding one testable integration claim.",
    ),
    "4.3 Human intervention architecture": (
        "study1",
        "Figure 3. Study 1 matched-budget benchmark and hard gates.",
        "Flow from frozen events to qualification, authorization, policy replay, and evaluation.",
    ),
    "5.4 Bounded human-correction replay": (
        "results",
        "Figure 4. Preliminary descriptive and technical evidence.",
        "Dashboard of stage signals, recorded review, C0 replay, and one correction.",
    ),
    "7. Work plan and milestones": (
        "programme",
        "Figure 5. Dependency of the three studies and integrated evaluation.",
        "Studies one through three feeding the gated integrated test.",
    ),
    "B. Gap Analysis": (
        "gap",
        "Figure B1. Established components versus the residual integration claim.",
        "Three research streams and a testable integrated gap.",
    ),
    "E. Preliminary Results": (
        "results",
        "Figure E1. Preliminary evidence with explicit claim boundary.",
        "Four aggregate preliminary results.",
    ),
    "H. RQ-to-study Traceability": (
        "programme",
        "Figure H1. Research-question ownership across the programme.",
        "Three studies feeding an integrated evaluation.",
    ),
}


def _render_markdown(
    doc: Document,
    source: Path,
    figures: dict[str, Path],
    *,
    skip_headings: set[str],
    start_at: str | None = None,
    major_page_breaks: bool = False,
    compact: bool = False,
) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    active = start_at is None
    for kind, value in _markdown_blocks(lines):
        if kind == "heading":
            level, text = value  # type: ignore[misc]
            if not active and text == start_at:
                active = True
            if not active or text in skip_headings:
                continue
            if major_page_breaks and level == 1 and len(doc.paragraphs) > 4:
                doc.add_page_break()
            paragraph = doc.add_heading(text, level=level)
            if text in FIGURE_MAP:
                key, caption, alt = FIGURE_MAP[text]
                _add_figure(doc, figures[key], caption, alt, width=7.0 if not compact else 6.8)
            continue
        if not active:
            continue
        if kind == "paragraph":
            text = value  # type: ignore[assignment]
            if text.startswith("**Status:") or text.startswith("**Evidence vocabulary:"):
                _add_status_box(doc, text)
                continue
            paragraph = doc.add_paragraph()
            _inline(
                paragraph,
                text,
                size=9.0
                if compact
                else (
                    8.3
                    if text.startswith(
                        (
                            "Akl,",
                            "Alfrink,",
                            "Ali,",
                            "Amershi,",
                            "Aroyo,",
                            "Bansal,",
                            "Bareinboim,",
                            "Ben-David,",
                            "Ben Chaaben,",
                            "Bian,",
                            "Boxwala,",
                            "Buçinca,",
                            "Chen,",
                            "Chow,",
                            "Cobbe,",
                            "Dasgupta,",
                            "de Kleer,",
                            "Doyle,",
                            "Galster,",
                            "Gebru,",
                            "Geifman,",
                            "Guo,",
                            "Hartvigsen,",
                            "Hevner,",
                            "Horvitz,",
                            "Hu,",
                            "Kapoor,",
                            "La Rosa,",
                            "Li,",
                            "Madaan,",
                            "Mao,",
                            "Mitchell,",
                            "Moreau,",
                            "Mozannar,",
                            "National Institute",
                            "Peffers,",
                            "Peleg,",
                            "Reinhartz-Berger,",
                            "Santoni",
                            "Schünemann,",
                            "Singh,",
                            "Smyth,",
                            "Thorn",
                            "Verma,",
                            "Wieringa,",
                            "Zhang,",
                        )
                    )
                    else 10.2
                ),
            )
            if text.startswith(
                (
                    "Akl,",
                    "Alfrink,",
                    "Ali,",
                    "Amershi,",
                    "Aroyo,",
                    "Bansal,",
                    "Bareinboim,",
                    "Ben-David,",
                    "Ben Chaaben,",
                    "Bian,",
                    "Boxwala,",
                    "Buçinca,",
                    "Chen,",
                    "Chow,",
                    "Cobbe,",
                    "Dasgupta,",
                    "de Kleer,",
                    "Doyle,",
                    "Galster,",
                    "Gebru,",
                    "Geifman,",
                    "Guo,",
                    "Hartvigsen,",
                    "Hevner,",
                    "Horvitz,",
                    "Hu,",
                    "Kapoor,",
                    "La Rosa,",
                    "Li,",
                    "Madaan,",
                    "Mao,",
                    "Mitchell,",
                    "Moreau,",
                    "Mozannar,",
                    "National Institute",
                    "Peffers,",
                    "Peleg,",
                    "Reinhartz-Berger,",
                    "Santoni",
                    "Schünemann,",
                    "Singh,",
                    "Smyth,",
                    "Thorn",
                    "Verma,",
                    "Wieringa,",
                    "Zhang,",
                )
            ):
                paragraph.paragraph_format.first_line_indent = Inches(-0.22)
                paragraph.paragraph_format.left_indent = Inches(0.22)
                paragraph.paragraph_format.space_after = Pt(2.4)
                paragraph.paragraph_format.line_spacing = 1.02
        elif kind in {"bullet", "number"}:
            paragraph = doc.add_paragraph(
                style="List Bullet" if kind == "bullet" else "List Number"
            )
            paragraph.paragraph_format.left_indent = Inches(0.26)
            paragraph.paragraph_format.first_line_indent = Inches(-0.16)
            paragraph.paragraph_format.space_after = Pt(2.5)
            _inline(paragraph, value, size=8.7 if compact else 9.8)  # type: ignore[arg-type]
        elif kind == "table":
            _add_table(doc, value, compact=compact)  # type: ignore[arg-type]
        elif kind == "code":
            _, code = value  # type: ignore[misc]
            for line in code.splitlines() or [""]:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.2)
                paragraph.paragraph_format.right_indent = Inches(0.1)
                paragraph.paragraph_format.space_after = Pt(0)
                _set_paragraph_shading(paragraph, "F1F5F8")
                run = paragraph.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(6.6 if compact else 7.5)
                run.font.color.rgb = _rgb(NAVY)


def _set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(_qname("w", "shd"))
    if shd is None:
        shd = OxmlElement(_xml_name("w", "shd"))
        p_pr.append(shd)
    shd.set(_qname("w", "fill"), fill)


def _static_contents(doc: Document, source: Path) -> None:
    headings: list[tuple[int, str]] = []
    for line in source.read_text(encoding="utf-8").splitlines():
        match = re.match(r"^(#{1,2}) (.+)$", line)
        if match and not match.group(2).startswith(("Governed Human", "Selective review")):
            headings.append((len(match.group(1)), match.group(2)))
    doc.add_heading("Contents", level=1)
    for level, text in headings:
        if level != 1:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(3)
        _inline(p, text, size=9.2, color=NAVY)
    doc.add_page_break()


def _save(doc: Document, path: Path, title: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.author = "Ali Hamed"
    doc.core_properties.subject = "VEGO-AI supervisor review package"
    doc.core_properties.keywords = "VEGO-AI, human judgment, escalation, proposal"
    update_fields = doc.settings.element.find(_qname("w", "updateFields"))
    if update_fields is None:
        update_fields = OxmlElement(_xml_name("w", "updateFields"))
        doc.settings.element.append(update_fields)
    update_fields.set(_qname("w", "val"), "true")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_proposal(figures: dict[str, Path], output: Path) -> Path:
    source = PACKAGE / "proposal-v2-candidate.md"
    doc = Document()
    _configure_document(doc)
    _title_page(
        doc,
        "Governed Human Judgment for Agentic Variability Exploration",
        "Selective review, traceable decisions, and controlled reuse in VEGO-AI",
        "Doctoral research proposal · Version 2 candidate",
    )
    _static_contents(doc, source)
    _render_markdown(
        doc,
        source,
        figures,
        skip_headings={
            "Governed Human Judgment for Agentic Variability Exploration",
            "Selective review, traceable decisions, and controlled reuse in VEGO-AI",
        },
        start_at="Abstract",
        major_page_breaks=True,
    )
    path = output / "VEGO_AI_Doctoral_Proposal_v2_Candidate_20260903.docx"
    _save(doc, path, "VEGO-AI Doctoral Proposal v2 Candidate")
    return path


def build_evidence(figures: dict[str, Path], output: Path) -> Path:
    source = PACKAGE / "supervisor-evidence-package.md"
    doc = Document()
    _configure_document(doc)
    _title_page(
        doc,
        "Supervisor Evidence Package",
        "Requirements · gap · protocol · results · comment closure · traceability · validation",
        "Study 1 and proposal revision evidence package",
    )
    _render_markdown(
        doc,
        source,
        figures,
        skip_headings={"VEGO-AI supervisor evidence package"},
        start_at="Executive answer",
        major_page_breaks=True,
    )
    path = output / "VEGO_AI_Supervisor_Evidence_Package_20260903.docx"
    _save(doc, path, "VEGO-AI Supervisor Evidence Package")
    return path


def build_one_page(figures: dict[str, Path], output: Path) -> Path:
    source = PACKAGE / "study1-one-page-plan.md"
    doc = Document()
    _configure_document(doc, compact=True)
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("VEGO-AI STUDY 1  ·  HUMAN ESCALATION BASELINE")
    run.font.name = "Aptos Display"
    run.font.size = Pt(17)
    run.font.bold = True
    run.font.color.rgb = _rgb(NAVY)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(3)
    _inline(
        sub,
        "One-page supervisor plan · 3 September 2026 · evidence-bound candidate",
        size=8.2,
        color=BLUE,
    )
    _render_markdown(
        doc,
        source,
        figures,
        skip_headings={"Study 1: where and when should VEGO-AI ask a human?"},
        start_at="Question and claim boundary",
        compact=True,
    )
    path = output / "VEGO_AI_Study1_One_Page_Plan_20260903.docx"
    _save(doc, path, "VEGO-AI Study 1 One-Page Plan")
    return path


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()
    output = args.output.resolve()
    figures = build_figures(output / "figures")
    paths = [
        build_proposal(figures, output),
        build_evidence(figures, output),
        build_one_page(figures, output),
    ]
    for path in paths:
        print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
