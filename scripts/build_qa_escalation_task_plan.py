"""Build the supervisor-facing Hebrew Q&A escalation task plan."""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from qa_task_plan_data import TASKS as PLAN_TASKS


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "VEGO-AI-qa-escalation-operational-task-plan-he.docx"

NAVY = "17365D"
BLUE = "2E74B5"
MUTED = "5B6573"
LIGHT = "EEF3F8"
PALE = "F7F9FB"
RED = "8B1E2D"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=55, start=85, bottom=55, end=85):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths=(1800, 7560)):
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_bidi(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run(run, size=8.4, bold=False, color="1F2937", font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:cs"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=8.7, bold=False, color="1F2937", before=0, after=2.5, align=WD_ALIGN_PARAGRAPH.RIGHT, style=None):
    p = doc.add_paragraph(style=style)
    set_bidi(p, align)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def add_task(doc, number, name, priority, fields):
    p = doc.add_paragraph()
    set_bidi(p)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"משימה {number} — {name}  [{priority}]")
    set_run(r, size=10.2, bold=True, color=NAVY)
    table = doc.add_table(rows=0, cols=2)
    labels = [
        "מטרה", "מה אני אבצע", "המקור לזיהוי האוטומטי", "ה-Dataset", "הפלט",
        "קריטריון השלמה", "מה נדרש ממני", "מה נדרש מאיריס וארנון", "מה חסר / מאתגר", "תלויות", "הערכת זמן",
    ]
    for label in labels:
        row = table.add_row()
        left, right = row.cells
        set_cell_width(left, 1800)
        set_cell_width(right, 7560)
        shade(left, LIGHT)
        shade(right, "FFFFFF")
        for cell in (left, right):
            set_cell_margins(cell)
        lp = left.paragraphs[0]
        set_bidi(lp)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(label)
        set_run(lr, size=7.6, bold=True, color=NAVY)
        rp = right.paragraphs[0]
        set_bidi(rp)
        rp.paragraph_format.space_after = Pt(0)
        rp.paragraph_format.line_spacing = 1.0
        rr = rp.add_run(fields[label])
        set_run(rr, size=7.55, color="263238")
    set_table_geometry(table)
    # restrained borders
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "3")
        node.set(qn("w:color"), "D9E2EC")


TASKS = [
    ("תיקון והקפאת מצאי תקשורת ה-Q&A", "P0 | הושלם טכנית; נדרש קיבוע תיעודי", {
        "מטרה": "לקבע בדיוק איזו ראיית תקשורת קיימת כיום ומה אינה זמינה.",
        "מה אני אבצע": "ראשית אחפש את interaction_log.jsonl המקורי בכל החבילות והמאגרים — צעד בעלות אפס לפני כל הרצה מחדש. לאחר מכן אאמת את 12 שאלות Agent 2 → Agent 1, את היעדר התשובות השמורות ואת המסלולים הנתמכים-בקוד אך לא נצפו; הסימון יהיה ANSWER_NOT_PERSISTED.",
        "המקור לזיהוי האוטומטי": "agentB_best_guidelines.json; eval_state.json; state.lang_qa_history/state.dom_qa_history; orchestrator; hash manifest; הגדרת interaction_log ב-eval_config.json וכתיבת response_raw ב-llm_client.py.",
        "ה-Dataset": "ucd_ch, ucd_pw, cd_ch, cd_pw; snapshot קפוא ובנפרד 30 רשומות snapshot של Agent B.",
        "הפלט": "דוח observability קפוא עם מטריצת תקשורת, ספירות, hash וגבול טענה; ולגבי ה-interaction log: זמינות, hash ומצאי שדות.",
        "קריטריון השלמה": "כל רשומה מזוהה ומקושרת ל-source hash; אין טענה שהתרחשה אי-מענה בזמן הריצה.",
        "מה נדרש ממני": "קיבוע הדוח והטרמינולוגיה במסמך נשלט.",
        "מה נדרש מאיריס וארנון": "רק אם לא אאתר את הקובץ בעצמי: לשלוח את interaction_log.jsonl המקורי אם הוא עדיין שמור. זו אינה בקשה לתיוג או לבדיקה ידנית.",
        "מה חסר / מאתגר": "ה-artifacts אינם משמרים את מלוא פרק התקשורת; ייתכן שנוצרה תשובה שלא יוצאה. גם אם ה-interaction log יימצא, הוא עשוי להחזיר פלטים גולמיים ו-provenance עשירים יותר אך אינו יכול להכיל תשובות יועץ שמעולם לא נוצרו; מאחר שה-evaluator לא הריץ לולאת מענה, אין לצפות לשחזור answer confidence משם.",
        "תלויות": "קבצי evaluation קפואים ו-hash manifest.",
        "הערכת זמן": "Ali: 0.5–1 שעה; מכונה: פחות מ-5 דקות.",
    }),
    ("הוספת observability מלאה ל-Q&A", "P0", {
        "מטרה": "לשמר בעתיד כל פרק Q&A שלם בלי לשנות את החלטות הסוכנים.",
        "מה אני אבצע": "אוסיף אירועי schema-versioned עם ID, סוכנים, שלב, שאלה, תשובה, confidence, evidence, episode, round, follow-up, convergence/termination ו-context.",
        "המקור לזיהוי האוטומטי": "נקודות היצירה וההשבה ב-orchestrator.py, qa_registry.py ו-state.py.",
        "ה-Dataset": "fixtures בטוחים והרצה מבוקרת מקומית; ללא העלאת טקסטי סטודנטים לענן.",
        "הפלט": "רשומות qa-communication-event-v1 חתומות ומקושרות ל-run manifest.",
        "קריטריון השלמה": "כל שדות החובה קיימים ו-baseline output אינו משתנה.",
        "מה נדרש ממני": "מימוש instrumentation, schema ובדיקות.",
        "מה נדרש מאיריס וארנון": "לא נדרש דבר, אלא אם שינוי runtime מחייב אישור governance.",
        "מה חסר / מאתגר": "שמירה מלאה מגדילה artifacts; יש להימנע מהשפעה על prompts ועל אי-דטרמיניזם.",
        "תלויות": "אישור repository לשינוי instrumentation ושדות runtime.",
        "הערכת זמן": "Ali: 4–6 שעות; בדיקות: 15–30 דקות.",
    }),
    ("אימות שה-instrumentation אינו משנה את המערכת", "P0", {
        "מטרה": "להוכיח שהתקשורת מתועדת נכון ושפלט VEGO-AI נשאר זהה.",
        "מה אני אבצע": "unit/integration tests, schema validation, mandatory-field coverage והשוואת hashes של baseline לפני ואחרי.",
        "המקור לזיהוי האוטומטי": "event schema, run manifests, canonical JSON hashes ו-baseline diff.",
        "ה-Dataset": "fixtures בטוחים והרצה מבוקרת מינימלית אחת.",
        "הפלט": "verification receipt עם pass/fail, coverage והשוואת hashes.",
        "קריטריון השלמה": "שתי ריצות זהות מחזירות artifacts קנוניים זהים וללא שדות חסרים.",
        "מה נדרש ממני": "כתיבת והרצת בדיקות ותיעוד חריגות.",
        "מה נדרש מאיריס וארנון": "לא נדרש דבר.",
        "מה חסר / מאתגר": "LLM runtime אינו דטרמיניסטי בהכרח; ההשוואה תיעשה על fixture/controlled run.",
        "תלויות": "משימה 2 וה-baseline הקפוא.",
        "הערכת זמן": "Ali: 2–3 שעות; ריצה: 15–45 דקות.",
    }),
    ("הרצת VEGO-AI עם observability מועשר", "P1", {
        "מטרה": "לאסוף corpus שמיש של פרקי Q&A בפועל.",
        "מה אני אבצע": "אבדוק inputs ואפעיל, במידת האפשר, Cheers ו-ParkWise בארבע ההגדרות; אדווח על הרצה מלאה/מופחתת ועל חסרים.",
        "המקור לזיהוי האוטומטי": "runtime logs, run manifest ורשומות qa-communication-event-v1.",
        "ה-Dataset": "Cheers ו-ParkWise, בכפוף להרשאות, קלטים ומשאבי runtime.",
        "הפלט": "corpus Q&A חתום, דוח כיסוי והודעת reproducibility.",
        "קריטריון השלמה": "כל setting מתועד עם input hash, episodes, זמני ריצה, שגיאות וגודל corpus; חסום מסומן כחסום.",
        "מה נדרש ממני": "בדיקת קלטים, dry-run, הרצה ותיעוד עלויות/זמנים.",
        "מה נדרש מאיריס וארנון": "רק אישור אם יש עלות API, החלטת פרסום או שינוי baseline.",
        "מה חסר / מאתגר": "קלטים, עלות API, אי-דטרמיניזם וזמן ריצה.",
        "תלויות": "משימות 2–3, inputs תקינים והרשאות runtime.",
        "הערכת זמן": "Ali: 2–4 שעות; מכונה/API: 1–6 שעות; עלות API תימדד.",
    }),
    ("חילוץ אוטומטי של אותות תקשורת", "P1", {
        "מטרה": "לגזור מן ה-corpus רק אותות escalation שניתנים למדידה בפועל.",
        "מה אני אבצע": "אחשב confidence, ANSWER_NOT_PERSISTED, evidence חסר, חזרה, follow-up, rounds, non-convergence, MAX_QA_ROUNDS וספירה גבוהה; mapping certainty ו-Agent 4 confidence יישארו הקשר.",
        "המקור לזיהוי האוטומטי": "שדות event v1, היסטוריית episode ו-runtime termination record.",
        "ה-Dataset": "corpus מועשר ממשימה 4; ה-snapshot הקיים ישמש sanity check בלבד.",
        "הפלט": "feature table קנונית עם זמינות, count ו-limitation לכל signal.",
        "קריטריון השלמה": "לכל feature יש rule, source field, availability ו-count; שום proxy אינו ground truth.",
        "מה נדרש ממני": "מימוש חישוב features ובדיקות determinism.",
        "מה נדרש מאיריס וארנון": "לא נדרש דבר.",
        "מה חסר / מאתגר": "אותות שאינם נשמרים יישארו unavailable ולא יומצאו.",
        "תלויות": "משימה 4 ושדות observability מלאים.",
        "הערכת זמן": "Ali: 2–3 שעות; ריצה: פחות מ-10 דקות.",
    }),
    ("בנייה והקפאה של Detector v1", "P1", {
        "מטרה": "להפיק ALERT או NO ALERT באופן שקוף, דטרמיניסטי וניתן להסבר.",
        "מה אני אבצע": "אגדיר rule-based OR עם reason codes, ללא machine learning, ללא tuning לפי labels וללא שינוי prompts או החלטות.",
        "המקור לזיהוי האוטומטי": "feature table ממשימה 5.",
        "ה-Dataset": "כל episodes שנאספו; אותה טבלת קלט לכל הפעלה.",
        "הפלט": "detector manifest, event-level decisions ו-reason-code inventory.",
        "קריטריון השלמה": "אותו קלט/גרסה/seed מפיק אותו decision hash; כל ALERT מקושר ל-feature ול-source event.",
        "מה נדרש ממני": "קביעת rules, מימוש, versioning ובדיקת reproducibility.",
        "מה נדרש מאיריס וארנון": "לכל היותר אישור מתודולוגי יחיד לכללי v1; אין בדיקת שורות.",
        "מה חסר / מאתגר": "ALERT הוא מועמד להתערבות, לא הוכחה שנדרשה התערבות.",
        "תלויות": "משימה 5; אין שימוש ב-labels אנושיים.",
        "הערכת זמן": "Ali: 2–3 שעות; ריצה: פחות מ-10 דקות.",
    }),
    ("ניתוח תיאורי אוטומטי", "P1", {
        "מטרה": "לדווח את מה שניתן למדוד ללא תיוג ידני.",
        "מה אני אבצע": "אחשב episodes, מספר/שיעור alerts, חלוקות agent/stage/target, signals, confidence, ANSWER_NOT_PERSISTED, rounds/follow-up ו-repeatability; אייצר טבלה ותרשים.",
        "המקור לזיהוי האוטומטי": "detector outputs, feature table ו-run manifests.",
        "ה-Dataset": "כל corpus Q&A המאומת ממשימה 4.",
        "הפלט": "descriptive-results JSON/CSV, תרשים ו-report קנוני.",
        "קריטריון השלמה": "כל מספר ניתן לשחזור מ-event IDs וה-hashes; אין true/false, precision, recall או accuracy.",
        "מה נדרש ממני": "הפקת הדוח, QA של מספרים ופרשנות תיאורית מוגבלת.",
        "מה נדרש מאיריס וארנון": "לא נדרש דבר.",
        "מה חסר / מאתגר": "ללא ground truth אי אפשר למדוד איכות זיהוי או missed intervention.",
        "תלויות": "משימות 5–6.",
        "הערכת זמן": "Ali: 2–3 שעות; ריצה: פחות מ-15 דקות.",
    }),
    ("קיבוע המגבלה והחלטה על ראיות עתידיות", "P2 | נדחה", {
        "מטרה": "לתעד את גבול הטענה ולהחליט אם תוצאה תיאורית מספיקה לאבן-הדרך המקדימה.",
        "מה אני אבצע": "אנסח limitation מפורש, אשמור reviewer sheets ככלי עתידי בלבד ולא אפעילם.",
        "המקור לזיהוי האוטומטי": "claim register, descriptive report ו-evidence boundary.",
        "ה-Dataset": "תוצרי משימות 1–7; אין dataset חדש.",
        "הפלט": "decision note עם סטטוס validation deferred.",
        "קריטריון השלמה": "התכנית אינה מבקשת מאיריס או מארנון לתייג או לבדוק אירועים.",
        "מה נדרש ממני": "שמירת גבול הטענה ועדכון התיעוד.",
        "מה נדרש מאיריס וארנון": "החלטה אחת: האם feasibility תיאורי מספיק, כאשר true/false validation נדחה.",
        "מה חסר / מאתגר": "אין ground truth אוטומטי לגיטימי; אין להמיר proxy לתווית מדעית.",
        "תלויות": "תוצאות משימה 7 והחלטת governance.",
        "הערכת זמן": "Ali: 0.5–1 שעה; אין runtime.",
    }),
]

# The shared plan data is the single source for current supervisor-facing task order.
TASKS = PLAN_TASKS


def add_summary_table(doc):
    headers = ["#", "משימה", "עדיפות", "מקור", "Dataset", "תוצר", "נדרש מאיריס/ארנון", "זמן"]
    rows = [
        ["1", "איתור interaction_log.jsonl", "P0", "eval_config / llm_client", "הרצה מקורית", "דוח זמינות/hash", "רק אם לא נמצא", "0.5–1 ש'"],
        ["2", "קיבוע מצב Q&A קיים", "P0", "snapshots / eval_state", "4 settings + 30 רשומות", "observability report", "לא נדרש", "0.5–1 ש'"],
        ["3", "הגדרת אירוע Q&A מלא", "P0", "state / schema", "fixtures", "event-v1 + validator", "לא נדרש", "1–2 ש'"],
        ["4", "instrumentation פסיבי", "P0", "orchestrator / writer", "fixtures", "אירועים", "לא נדרש", "4–6 ש'"],
        ["5", "אימות instrumentation", "P0", "hashes / tests", "fixture", "receipt", "לא נדרש", "2–3 ש'"],
        ["6", "הרצה מבוקרת של setting אחד", "P1", "runtime logs", "setting יחיד", "corpus + cost", "רק אם יש עלות", "2–4 ש'"],
        ["7", "חילוץ אותות", "P1", "event features", "corpus", "feature table", "לא נדרש", "2–3 ש'"],
        ["8", "detector + descriptive analysis", "P1", "features / manifest", "corpus", "decisions + report", "לא נדרש", "2–3 ש'"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    widths = (360, 1500, 520, 1380, 1250, 1500, 1800, 1050)
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        p = cell.paragraphs[0]
        set_bidi(p, WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, size=6.7, bold=True, color="FFFFFF")
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            shade(cell, "F7F9FB" if len(table.rows) % 2 == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            set_bidi(p, WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 7) else WD_ALIGN_PARAGRAPH.RIGHT)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run(r, size=6.45, bold=False, color="263238")
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "3")
        node.set(qn("w:color"), "D9E2EC")
    return table


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.05)
    sec.bottom_margin = Cm(0.9)
    sec.left_margin = Cm(1.0)
    sec.right_margin = Cm(1.0)
    sec.header_distance = Cm(0.45)
    sec.footer_distance = Cm(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(8.7)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.02

    header = sec.header.paragraphs[0]
    set_bidi(header, WD_ALIGN_PARAGRAPH.RIGHT)
    header.paragraph_format.space_after = Pt(0)
    hr = header.add_run("VEGO-AI  |  תכנית עבודה אופרטיבית  |  03.09.2026")
    set_run(hr, size=7.2, bold=True, color=MUTED)
    footer = sec.footer.paragraphs[0]
    set_bidi(footer, WD_ALIGN_PARAGRAPH.CENTER)
    fr = footer.add_run("מסמך עבודה לבחינה אנושית — אין בו תיוג או אישור מדעי")
    set_run(fr, size=7.1, color=MUTED)

    add_para(doc, "תכנית עבודה אופרטיבית", size=18, bold=True, color=NAVY, after=1.5)
    add_para(doc, "זיהוי אוטומטי של צורך בהתערבות אנושית מתוך תקשורת Q&A ב-VEGO-AI", size=11.5, bold=True, color=BLUE, after=3)
    add_para(doc, "לפרופ' Iris Reinhartz-Berger ולפרופ' Arnon Sturm  |  סטטוס: תכנית לביצוע ללא תיוג או בדיקה ידנית בשלב הנוכחי", size=8.2, color=MUTED, after=5)
    p = add_para(doc, "בהתאם להנחיה שלא יבוצעו תיוגים או בדיקות ידניות, התכנית מתמקדת בשלב זה באיסוף מלא של תקשורת ה-Q&A, בזיהוי אוטומטי של התראות ובניתוח תיאורי. בדיקת נכונות/שגיאות ההתראות תידחה לשלב שבו יהיה מקור תיוג עצמאי.", size=9.1, bold=True, color=NAVY, after=4)
    shade_box = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    shade_box.append(p_bdr)
    add_para(doc, "מצב Q&A מאומת: ה-snapshot הקפוא מכיל 12 שאלות Agent 2 → Agent 1, ללא תשובה תואמת שנשמרה, ללא answer confidence/evidence וללא מידע בר-שחזור על follow-up, rounds או convergence. המונח המחייב הוא ANSWER_NOT_PERSISTED; אין להסיק מהתרחשות בזמן הריצה.", size=8.0, color="263238", after=3)
    add_para(doc, "בהיעדר מקור ייחוס עצמאי וללא תיוג ידני בשלב הנוכחי, ניתן למדוד באופן מלא את מספר ההתראות, התפלגותן והאותות שהפעילו אותן, אך לא ניתן לקבוע באופן אמפירי בשלב זה אילו מהן True/False.", size=8.0, color=RED, bold=True, after=3)
    add_para(doc, "רמות עדיפות: P0 — תנאי אפשרות; P1 — תוצאה תיאורית מקדימה; P2 — אימות עתידי שנדחה.", size=7.8, color=MUTED, after=3)
    add_summary_table(doc)
    add_para(doc, "פירוט המשימות", size=9.2, bold=True, color=NAVY, before=4, after=1)
    for i, (name, priority, fields) in enumerate(TASKS, 1):
        add_task(doc, i, name, priority, fields)
    add_para(doc, "סה\"כ עבודת Ali נטו: כ-14–23 שעות לכל שמונה המשימות. זמן ריצה/API: כ-1–6 שעות להרצה מבוקרת אחת, בתוספת פחות משעה לבדיקות ולניתוח מקומי. זמן חסום/המתנה תלוי ב-interaction_log, inputs, הרשאות ועלות אפשרית ואינו נכלל בזמן העבודה.", size=8.2, bold=True, color=NAVY, before=5, after=2)
    add_para(doc, "מה נדרש מאיריס וארנון: רק להעביר interaction_log.jsonl אם הוא קיים, או לאשר controlled run אחד ועלות API אם שחזור הראיות נכשל. אין בקשה לתייג, לקרוא שורות או לאשר כללי detector.", size=8.3, bold=True, color=RED, after=0)
    doc.core_properties.title = "תכנית עבודה אופרטיבית — Q&A escalation ב-VEGO-AI"
    doc.core_properties.subject = "Supervisor-facing operational task list"
    doc.core_properties.author = "VEGO-AI Research"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
