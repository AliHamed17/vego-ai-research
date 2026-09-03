from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
SOURCE = (
    ROOT
    / "docs"
    / "research"
    / "phd-proposal"
    / "2026-09-03-preliminary-human-intervention-experiment.en.md"
)
BUILDER = ROOT / "scripts" / "build_paper.py"


def test_supervisor_one_pager_has_the_required_evidence_bound_design(tmp_path: Path) -> None:
    output = tmp_path / "one-page.docx"
    result = subprocess.run(
        [
            sys.executable,
            str(BUILDER),
            str(SOURCE),
            str(output),
            "--figures",
            str(ROOT),
        ],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )

    assert result.returncode == 0, result.stderr
    document = Document(output)
    text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    )

    required_sections = [
        "1. Objective",
        "2. Research Question / Sub-question",
        "3. Experimental Cases",
        "4. Baseline",
        "5. Human-Intervention Trigger",
        "6. Human Intervention",
        "7. Human-Assisted Condition",
        "8. Comparison and Evaluation",
        "9. Expected Preliminary Contribution",
        "10. Limitations",
    ]
    for section in required_sections:
        assert section in text

    assert "Condition A - Autonomous VEGO-AI" in text
    assert "Condition B - Human-Assisted VEGO-AI" in text
    assert (
        "The human intervention is simulated/controlled for this preliminary "
        "feasibility experiment and does not constitute a human-subject user study."
    ) in text
    assert "When and how, in variability exploration scenarios" in text
    assert "currently automatically measurable" in text
    assert "manually identified for this preliminary experiment" in text
    assert "credible reference" in text
    assert "To be measured" in text
    assert "Demonstrate feasibility, not prove effectiveness." in text
    assert "No intervention outcome is claimed before independent evaluation." in text
    assert "17.5/27" not in text and "16.5/27" not in text

    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 4
    assert [cell.text for cell in table.rows[0].cells] == [
        "Case",
        "Baseline Issue",
        "Trigger",
        "Human Input",
        "Result After Intervention",
        "Reference",
        "Outcome",
    ]
    widths = [int(cell._tc.tcPr.tcW.w) for cell in table.rows[0].cells]
    assert widths[0] < min(widths[1:])
    assert widths[0] >= 650

    section = document.sections[0]
    assert round(section.page_width.cm, 1) == 21.0
    assert round(section.page_height.cm, 1) == 29.7
