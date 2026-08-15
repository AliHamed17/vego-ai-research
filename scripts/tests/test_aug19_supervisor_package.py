from __future__ import annotations

import importlib.util
import json
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
SCRIPT = ROOT / "scripts/validate_aug19_supervisor_package.py"
SPEC = importlib.util.spec_from_file_location("validate_aug19_supervisor_package", SCRIPT)
assert SPEC and SPEC.loader
MODULE = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = MODULE
SPEC.loader.exec_module(MODULE)

VISUAL_SCRIPT = ROOT / "scripts/record_aug19_visual_inspection.py"
BUILDER_SCRIPT = ROOT / "scripts/build_aug19_supervisor_package.py"

PACKAGE_ROOT = ROOT / "docs/research/meetings/2026-08-19-supervisor-package"
FINAL_DIR = PACKAGE_ROOT / "final"
CONTENT = PACKAGE_ROOT / "source/package-content.json"


EXPECTED_FILES = {
    "Chapter_2_Literature_Review_EN.docx",
    "Chapter_2_Literature_Review_EN.pdf",
    "Chapter_2_Literature_Review_HE.docx",
    "Chapter_2_Literature_Review_HE.pdf",
    "Human_Agent_Classification_Bilingual.pptx",
    "Human_Agent_Classification_Bilingual.pdf",
    "Supervisor_PreRead_EN.pdf",
    "Supervisor_PreRead_HE.pdf",
    "Supervisor_Tracker_and_Decisions_Bilingual.pdf",
    "package-manifest.sanitized.json",
}


def _minimal_content() -> dict[str, object]:
    return {
        "meeting": {
            "date": "2026-08-19",
            "time": "09:00-10:00",
            "timezone": "Asia/Jerusalem",
        },
        "boundaries": {
            "scope": "Literature only",
            "section_4": "Frozen pre-existing methodology draft",
            "rq_wording": "Open",
            "e6_exploration_identification": "Open",
            "e8_human_expert": "Open",
            "plan_a_b": "Open",
            "methodology_conflict": "Open",
            "exp_005": "0/24",
            "medical_readiness": "0/6",
            "human_screening": "Pending",
        },
        "chapter2": {
            "en": [{"id": "C2-01", "title": "Scope", "paragraphs": ["Evidence-bounded."]}],
            "he": [{"id": "C2-01", "title": "תחום", "paragraphs": ["מוגבל לראיות."]}],
        },
        "decisions": [
            {
                "canonical_control_id": "D12-001",
                "state": "Open",
                "question_en": "Retain?",
                "question_he": "להשאיר?",
            }
        ],
        "workbook": {
            "title": "VEGO-AI PhD Literature Workbook v0.1",
            "logical_alias": "VEGO-AI-PHD-LITERATURE-WORKBOOK",
            "delivery": (
                "Native workbook referenced by logical alias; not duplicated; "
                "live link withheld until Ali-approved release"
            ),
            "link_status": "Withheld pending Ali-approved private release binding",
            "private_binding_policy": (
                "Inject the live URL only from an ignored private binding after Ali "
                "approves the exact release package"
            ),
        },
    }


class August19PackageTests(unittest.TestCase):
    def test_pypdf_security_gate_rejects_runtime_or_receipt_below_6_15(self) -> None:
        for installed, receipt in (
            ("6.14.2", "6.15.0"),
            ("6.15.0", "6.10.0"),
            ("6.15.0rc1", "6.15.0"),
            ("not-a-version", "6.15.0"),
        ):
            with self.subTest(installed=installed, receipt=receipt), self.assertRaisesRegex(
                MODULE.PackageValidationError, "pypdf.*6.15.0"
            ):
                MODULE.validate_pypdf_runtime(installed, receipt)

        MODULE.validate_pypdf_runtime("6.15.0", "6.15.0")
        MODULE.validate_pypdf_runtime("7.0.0", "6.15.1")

    def test_required_output_contract_is_exactly_ten_files(self) -> None:
        self.assertEqual(set(MODULE.REQUIRED_FILES), EXPECTED_FILES)
        self.assertEqual(len(MODULE.REQUIRED_FILES), 10)

    def test_package_validator_rejects_missing_and_extra_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            package = Path(tmp)
            for name in EXPECTED_FILES - {"Supervisor_PreRead_HE.pdf"}:
                (package / name).write_bytes(b"fixture")
            with self.assertRaisesRegex(MODULE.PackageValidationError, "missing"):
                MODULE.validate_file_set(package)

            (package / "Supervisor_PreRead_HE.pdf").write_bytes(b"fixture")
            (package / "private-notes.txt").write_text("must not ship", encoding="utf-8")
            with self.assertRaisesRegex(MODULE.PackageValidationError, "extra"):
                MODULE.validate_file_set(package)

    def test_content_contract_requires_open_decisions_and_evidence_boundaries(self) -> None:
        content = _minimal_content()
        MODULE.validate_content_contract(content)

        content["boundaries"]["rq_wording"] = "Approved"  # type: ignore[index]
        with self.assertRaisesRegex(MODULE.PackageValidationError, "rq_wording"):
            MODULE.validate_content_contract(content)

        content = _minimal_content()
        content["boundaries"]["exp_005"] = "24/24"  # type: ignore[index]
        with self.assertRaisesRegex(MODULE.PackageValidationError, "exp_005"):
            MODULE.validate_content_contract(content)

    def test_content_contract_requires_english_hebrew_section_id_parity(self) -> None:
        content = _minimal_content()
        content["chapter2"]["he"][0]["id"] = "C2-99"  # type: ignore[index]
        with self.assertRaisesRegex(MODULE.PackageValidationError, "section IDs"):
            MODULE.validate_content_contract(content)

    def test_canonical_status_records_build_complete_but_release_pending(self) -> None:
        content = json.loads(CONTENT.read_text(encoding="utf-8"))

        self.assertEqual(
            content["status"],
            (
                "Canonical provisional controls mapped; local package build complete; "
                "release and delivery pending Ali approval"
            ),
        )
        self.assertNotIn("build pending", content["status"].lower())

    def test_content_contract_rejects_pending_canonical_control_mapping(self) -> None:
        content = _minimal_content()
        content["decisions"][0]["canonical_control_id"] = "Pending media-control handoff"  # type: ignore[index]
        with self.assertRaisesRegex(MODULE.PackageValidationError, "canonical control"):
            MODULE.validate_content_contract(content)

    def test_hebrew_overclaim_patterns_fail_closed(self) -> None:
        findings = MODULE.scan_forbidden_claims(
            {"hebrew.pdf": "שאלות המחקר אושרו והמתודולוגיה אושרה."}
        )
        self.assertEqual(len(findings), 2)

    def test_manifest_artifacts_reject_path_traversal_duplicate_and_missing_entries(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            package = Path(tmp)
            for name in EXPECTED_FILES - {"package-manifest.sanitized.json"}:
                (package / name).write_bytes(name.encode("utf-8"))

            rows = [
                {"file": name, "sha256": MODULE.sha256_file(package / name)}
                for name in sorted(EXPECTED_FILES - {"package-manifest.sanitized.json"})
            ]
            self.assertEqual(MODULE.validate_manifest_artifacts(rows, package), 9)

            traversal = [dict(row) for row in rows]
            traversal[0]["file"] = "../private.txt"
            with self.assertRaisesRegex(MODULE.PackageValidationError, "unsafe artifact path"):
                MODULE.validate_manifest_artifacts(traversal, package)

            duplicate = [dict(row) for row in rows]
            duplicate[-1] = dict(duplicate[0])
            with self.assertRaisesRegex(MODULE.PackageValidationError, "duplicate"):
                MODULE.validate_manifest_artifacts(duplicate, package)

            with self.assertRaisesRegex(MODULE.PackageValidationError, "artifact set"):
                MODULE.validate_manifest_artifacts(rows[:-1], package)

    def test_manifest_bindings_require_exact_sources_runtimes_and_render_receipts(self) -> None:
        expected = {
            "acl_corpus_git_commit": "f2bcd796fad33dc578983108cf29ef4cd305a4a1",
            "taxonomy_repository_commit": "7b3ba9deefe99172748582f6025d995ccc2a6f86",
            "python_runtime": "bundled-python",
            "node_runtime": "bundled-node",
            "build_receipt_sha256": "a" * 64,
            "render_receipt_sha256": "b" * 64,
        }
        manifest = {
            "bindings": dict(expected),
        }
        MODULE.validate_manifest_bindings(manifest, expected)

        manifest["bindings"]["acl_corpus_git_commit"] = "stale"
        with self.assertRaisesRegex(MODULE.PackageValidationError, "binding"):
            MODULE.validate_manifest_bindings(manifest, expected)

    def test_pptx_notes_require_sources_block_and_exact_urls(self) -> None:
        required = {
            "https://aclanthology.org/2026.findings-acl.1811/",
            "https://github.com/HenryPengZou/Awesome-Human-Agent-Collaboration-Interaction-Systems/tree/7b3ba9deefe99172748582f6025d995ccc2a6f86",
        }
        with tempfile.TemporaryDirectory() as tmp:
            pptx = Path(tmp) / "slide.pptx"
            with zipfile.ZipFile(pptx, "w") as archive:
                archive.writestr(
                    "ppt/notesSlides/notesSlide1.xml",
                    "<p:notes><a:t>[Sources]</a:t>"
                    + "".join(f"<a:t>{url}</a:t>" for url in sorted(required))
                    + "</p:notes>",
                )
            MODULE.validate_pptx_source_notes(pptx, required)

            with zipfile.ZipFile(pptx, "w") as archive:
                archive.writestr("ppt/notesSlides/notesSlide1.xml", "<p:notes><a:t>No sources</a:t></p:notes>")
            with self.assertRaisesRegex(MODULE.PackageValidationError, r"\[Sources\]"):
                MODULE.validate_pptx_source_notes(pptx, required)

    def test_hebrew_research_terminology_rejects_provenance_as_originality(self) -> None:
        content = _minimal_content()
        content["chapter2"]["en"][0]["paragraphs"] = [  # type: ignore[index]
            "Provenance and accountable authority must be preserved."
        ]
        content["chapter2"]["he"][0]["paragraphs"] = [  # type: ignore[index]
            "מקוריות וסמכות אחראית חייבות להישמר."
        ]
        with self.assertRaisesRegex(MODULE.PackageValidationError, "provenance"):
            MODULE.validate_hebrew_terminology(content)

    def test_hebrew_research_terminology_rejects_unnatural_agentic_phrases(self) -> None:
        for phrase in ("מוגבלת-ראיות", "סוכנית"):
            content = _minimal_content()
            content["chapter2"]["he"][0]["paragraphs"] = [phrase]  # type: ignore[index]
            with self.subTest(phrase=phrase), self.assertRaisesRegex(
                MODULE.PackageValidationError, "Hebrew terminology"
            ):
                MODULE.validate_hebrew_terminology(content)

    def test_docx_requires_hyperlinks_and_hebrew_bidi_tokens(self) -> None:
        urls = {
            "https://aclanthology.org/2026.findings-acl.1811/",
            "https://example.org/workbook",
        }
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "he.docx"
            rels = "".join(
                f'<Relationship Id="r{index}" Target="{url}" TargetMode="External"/>'
                for index, url in enumerate(sorted(urls), start=1)
            )
            with zipfile.ZipFile(docx, "w") as archive:
                archive.writestr(
                    "word/document.xml",
                    '<w:document xmlns:w="w"><w:p><w:pPr><w:bidi/></w:pPr>'
                    '<w:r><w:rPr><w:rtl/></w:rPr><w:t>עברית</w:t></w:r></w:p></w:document>',
                )
                archive.writestr(
                    "word/_rels/document.xml.rels",
                    f'<Relationships xmlns="r">{rels}</Relationships>',
                )
            MODULE.validate_docx_hyperlinks(docx, urls)
            MODULE.validate_docx_rtl(docx, minimum_bidi=1, minimum_rtl=1)

            with self.assertRaisesRegex(MODULE.PackageValidationError, "hyperlinks"):
                MODULE.validate_docx_hyperlinks(docx, urls | {"https://missing.example"})

    def test_pdf_requires_clickable_uri_annotations(self) -> None:
        from pypdf import PdfWriter

        with tempfile.TemporaryDirectory() as tmp:
            pdf = Path(tmp) / "links.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=612, height=792)
            writer.add_uri(0, "https://example.org/source", (10, 10, 100, 30))
            with pdf.open("wb") as handle:
                writer.write(handle)
            MODULE.validate_pdf_hyperlinks(pdf, {"https://example.org/source"})
            with self.assertRaisesRegex(MODULE.PackageValidationError, "PDF hyperlinks"):
                MODULE.validate_pdf_hyperlinks(pdf, {"https://example.org/missing"})

    def test_visible_source_parity_requires_control_and_evidence_tokens(self) -> None:
        required = {"C2-01", "C2-ACL-01", "0/116", "0/24", "0/6"}
        MODULE.validate_visible_tokens(
            {"chapter.docx": "C2-01 C2-ACL-01 0/116 0/24 0/6"}, required
        )
        MODULE.validate_visible_tokens(
            {
                "chapter-he.docx": (
                    "\u200eC2\u201101\u200e \u200eC2\u2011ACL\u201101\u200e 0/116 0/24 0/6"
                )
            },
            required,
        )
        with self.assertRaisesRegex(MODULE.PackageValidationError, "visible source parity"):
            MODULE.validate_visible_tokens({"chapter.pdf": "C2-01 0/24"}, required)

    def test_visible_source_parity_tolerates_pdf_hyphen_extraction_artifacts(self) -> None:
        MODULE.validate_visible_tokens(
            {
                "chapter.pdf": (
                    "Native workbook [VEGO -AI-PHD-LITERATURE-\nWORKBOOK] | link withheld"
                )
            },
            {"VEGO-AI-PHD-LITERATURE-WORKBOOK"},
        )

    def test_content_contract_requires_opaque_workbook_alias_and_withheld_link(self) -> None:
        content = _minimal_content()
        MODULE.validate_content_contract(content)

        content["workbook"]["url"] = (  # type: ignore[index]
            "https://docs.google.com/"
            + "spreadsheets/d/"
            + "1SyntheticRestrictedWorkbookId"
        )
        with self.assertRaisesRegex(MODULE.PackageValidationError, "workbook.*restricted"):
            MODULE.validate_content_contract(content)

    def test_restricted_workbook_reference_scan_covers_text_and_relationships(self) -> None:
        safe = {"review.txt": "VEGO-AI-PHD-LITERATURE-WORKBOOK - link withheld"}
        MODULE.validate_restricted_workbook_references(safe)

        leaked = {
            "chapter.docx": (
                '<Relationship Target="https://docs.google.com/'
                + "spreadsheets/d/"
                + '1SyntheticRestrictedWorkbookId"/>'
            )
        }
        with self.assertRaisesRegex(MODULE.PackageValidationError, "restricted workbook"):
            MODULE.validate_restricted_workbook_references(leaked)

    def test_tracked_package_source_withholds_live_workbook_resource_id(self) -> None:
        package_text = CONTENT.read_text(encoding="utf-8")
        self.assertNotIn("docs.google.com/" + "spreadsheets/d/", package_text)
        self.assertIn("VEGO-AI-PHD-LITERATURE-WORKBOOK", package_text)

    def test_private_workbook_binding_requires_explicit_approval_and_private_file(self) -> None:
        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_private_binding", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)

        content = _minimal_content()
        with self.assertRaisesRegex(RuntimeError, "requires a private workbook binding"):
            builder.load_private_workbook_binding(
                content, None, ali_approved_release=True
            )

        with tempfile.TemporaryDirectory() as tmp:
            binding = Path(tmp) / "workbook-binding.private.json"
            url = (
                "https://docs.google.com/"
                + "spreadsheets/d/"
                + "1SyntheticRestrictedWorkbookId"
            )
            binding.write_text(
                json.dumps(
                    {
                        "logical_alias": "VEGO-AI-PHD-LITERATURE-WORKBOOK",
                        "url": url,
                    }
                ),
                encoding="utf-8",
            )
            with self.assertRaisesRegex(RuntimeError, "explicit Ali-approved release"):
                builder.load_private_workbook_binding(
                    content, binding, ali_approved_release=False
                )
            resolved_url, binding_hash = builder.load_private_workbook_binding(
                content, binding, ali_approved_release=True
            )
            self.assertEqual(resolved_url, url)
            self.assertRegex(binding_hash or "", r"^[0-9a-f]{64}$")

    def test_tracker_has_controlled_bilingual_title_and_decision_recording_fields(self) -> None:
        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_tracker", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)

        content = json.loads(CONTENT.read_text(encoding="utf-8"))
        with tempfile.TemporaryDirectory() as tmp:
            tracker_path = Path(tmp) / "tracker.docx"
            builder.build_tracker(content, tracker_path)
            tracker = Document(tracker_path)

            paragraph_text = [paragraph.text for paragraph in tracker.paragraphs]
            self.assertIn("Supervisor tracker and decisions", paragraph_text)
            self.assertIn("מעקב והחלטות למנחים", paragraph_text)
            self.assertNotIn(
                "Supervisor tracker and decisions | מעקב והחלטות למנחים",
                paragraph_text,
            )

            expected_headers = {
                "Canonical control ID",
                "Selected outcome | תוצאה שנבחרה",
                "Correction / read-back | תיקון / הקראה חוזרת",
                "Approver / timestamp | מאשר/ת / חותמת זמן",
                "Owner / due date | אחראי/ת / מועד יעד",
                "Evidence link | קישור לראיה",
            }
            recording_tables = [
                table
                for table in tracker.tables
                if {cell.text for cell in table.rows[0].cells} == expected_headers
            ]
            self.assertEqual(len(recording_tables), 1)
            recording_table = recording_tables[0]
            self.assertEqual(len(recording_table.rows), len(content["decisions"]) + 1)
            for row in recording_table.rows[1:]:
                self.assertEqual(len(row.cells), 6)
                self.assertTrue(all(cell.text for cell in row.cells))
                self.assertTrue(
                    all("Pending | ממתין" in cell.text for cell in row.cells[1:])
                )

    def test_tracker_control_ids_use_nonbreaking_display_hyphens(self) -> None:
        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_control_ids", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)

        display = builder.protect_canonical_control_ids(
            "D12-002; Q12-009; R12-007"
        )
        self.assertEqual(display, "D12‑002; Q12‑009; R12‑007")
        self.assertNotIn("D12-002", display)

    def test_office_export_is_bounded_atomic_and_pid_scoped(self) -> None:
        orchestrator = (ROOT / "scripts/export_aug19_office_artifacts.ps1").read_text(
            encoding="utf-8"
        )
        worker_path = ROOT / "scripts/export_aug19_office_worker.ps1"
        self.assertTrue(worker_path.is_file(), "per-file Office worker is required")
        worker = worker_path.read_text(encoding="utf-8")

        for required in (
            "Start-Process",
            "-WindowStyle Hidden",
            "WaitForExit",
            "TimeoutSeconds",
            "MaxAttempts = 2",
            "automation-pid",
            "File]::Replace",
            "File]::Move",
        ):
            self.assertIn(required, orchestrator)
        self.assertNotRegex(orchestrator, r"Get-Process\s+(?:WINWORD|POWERPNT).+Stop-Process")

        self.assertIn("Microsoft Print to PDF", worker)
        self.assertIn("UpdateFieldsAtPrint = $false", worker)
        self.assertIn("DisplayAlerts = 0", worker)
        self.assertIn("Resolve-AutomationPid", worker)
        self.assertIn("BaselinePids", worker)
        self.assertIn("exactly one new", worker)
        self.assertNotIn(".Fields.Update", worker)
        self.assertNotIn(".Visible = -1", worker)
        self.assertNotIn("PowerPoint.Application", worker)
        self.assertNotIn("POWERPNT", orchestrator)

    def test_slide_pdf_is_built_from_artifact_tool_render_without_powerpoint_com(self) -> None:
        spec = importlib.util.spec_from_file_location("build_aug19_supervisor_package", BUILDER_SCRIPT)
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        from PIL import Image

        with tempfile.TemporaryDirectory() as tmp:
            folder = Path(tmp)
            png = folder / "slide.png"
            pdf = folder / "slide.pdf"
            Image.new("RGB", (1280, 720), "white").save(png)
            builder.create_slide_pdf(png, pdf)
            reader = __import__("pypdf").PdfReader(str(pdf))
            self.assertEqual(len(reader.pages), 1)
            self.assertEqual(tuple(round(float(value)) for value in reader.pages[0].mediabox[2:]), (960, 540))

    def test_slide_pdf_builder_uses_locked_pillow_not_unlocked_reportlab(self) -> None:
        builder_source = BUILDER_SCRIPT.read_text(encoding="utf-8")
        self.assertNotIn("reportlab", builder_source)
        self.assertIn("from PIL import Image", builder_source)

    def test_tracked_json_writers_emit_canonical_lf_bytes(self) -> None:
        builder_spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_lf", BUILDER_SCRIPT
        )
        assert builder_spec and builder_spec.loader
        builder = importlib.util.module_from_spec(builder_spec)
        sys.modules[builder_spec.name] = builder
        builder_spec.loader.exec_module(builder)

        visual_spec = importlib.util.spec_from_file_location(
            "record_aug19_visual_inspection_lf", VISUAL_SCRIPT
        )
        assert visual_spec and visual_spec.loader
        visual = importlib.util.module_from_spec(visual_spec)
        sys.modules[visual_spec.name] = visual
        visual_spec.loader.exec_module(visual)

        with tempfile.TemporaryDirectory() as tmp:
            for index, writer in enumerate((builder.write_text_lf, visual.write_text_lf), start=1):
                output = Path(tmp) / f"receipt-{index}.json"
                writer(output, "{\n  \"status\": \"ready\"\n}\n")
                payload = output.read_bytes()
                self.assertNotIn(b"\r\n", payload)
                self.assertEqual(payload.count(b"\n"), 3)

    def test_office_export_uses_current_host_and_quotes_process_arguments(self) -> None:
        orchestrator = (ROOT / "scripts/export_aug19_office_artifacts.ps1").read_text(
            encoding="utf-8"
        )
        self.assertIn("(Get-Process -Id $PID).Path", orchestrator)
        self.assertIn("Test-Path -LiteralPath $powershellExe", orchestrator)
        self.assertIn("ConvertTo-ProcessArgument", orchestrator)
        self.assertNotIn("Join-Path $PSHOME 'powershell.exe'", orchestrator)

    def test_tracked_package_sources_contain_no_user_profile_paths_or_frozen_acl_commit(self) -> None:
        paths = (
            ROOT / "scripts/build_aug19_supervisor_package.py",
            ROOT / "scripts/build_aug19_classification_slide.mjs",
            ROOT / "scripts/export_aug19_office_artifacts.ps1",
            ROOT / "scripts/export_aug19_office_worker.ps1",
            ROOT / "scripts/validate_aug19_supervisor_package.py",
        )
        for path in paths:
            text = path.read_text(encoding="utf-8")
            with self.subTest(path=path.name):
                self.assertNotRegex(text, r"[A-Za-z]:\\Users\\[^\\\r\n]+")
                self.assertNotRegex(text, r"[A-Za-z]:\\\\Users\\\\[^\\\r\n]+")
        builder = paths[0].read_text(encoding="utf-8")
        self.assertNotIn(
            'ACL_CORPUS_GIT_COMMIT = "f2bcd796fad33dc578983108cf29ef4cd305a4a1"',
            builder,
        )
        self.assertIn("--acl-corpus-commit", builder)
        self.assertIn("--powershell-exe", builder)
        self.assertNotIn('["powershell.exe",', builder)

    def test_visual_receipt_validator_rejects_missing_or_changed_render(self) -> None:
        self.assertTrue(VISUAL_SCRIPT.is_file(), "visual receipt recorder is required")
        spec = importlib.util.spec_from_file_location("record_aug19_visual_inspection", VISUAL_SCRIPT)
        assert spec and spec.loader
        visual = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = visual
        spec.loader.exec_module(visual)
        with tempfile.TemporaryDirectory() as tmp:
            render_dir = Path(tmp)
            rendered = render_dir / "page-1.png"
            rendered.write_bytes(b"rendered")
            index = {
                "items": [
                    {
                        "artifact": "chapter.pdf",
                        "page_or_slide": 1,
                        "render_file": rendered.name,
                        "sha256": MODULE.sha256_file(rendered),
                    }
                ]
            }
            self.assertEqual(visual.validate_render_index(index, render_dir), 1)
            rendered.write_bytes(b"changed")
            with self.assertRaisesRegex(ValueError, "render hash"):
                visual.validate_render_index(index, render_dir)

    def test_renderer_resolution_accepts_bundled_cmd_wrapper(self) -> None:
        spec = importlib.util.spec_from_file_location("build_aug19_supervisor_package", BUILDER_SCRIPT)
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        with tempfile.TemporaryDirectory() as tmp:
            runtime_dir = Path(tmp)
            wrapper = runtime_dir / "pdftoppm.cmd"
            wrapper.write_text("@echo off\r\n", encoding="utf-8")
            self.assertEqual(builder.resolve_pdftoppm(runtime_dir), wrapper.resolve())

    def test_renderer_resolution_prefers_the_bundled_native_executable(self) -> None:
        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_native", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        with tempfile.TemporaryDirectory() as tmp:
            dependency_root = Path(tmp)
            override = dependency_root / "bin/override"
            native = dependency_root / "native/poppler/Library/bin/pdftoppm.exe"
            override.mkdir(parents=True)
            native.parent.mkdir(parents=True)
            (override / "pdftoppm.cmd").write_text("@exit /b 3\r\n", encoding="utf-8")
            native.write_bytes(b"bundled-native-renderer")
            self.assertEqual(builder.resolve_pdftoppm(override), native.resolve())

    def test_render_page_filter_excludes_the_artifact_tool_preview(self) -> None:
        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_page_filter", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        with tempfile.TemporaryDirectory() as tmp:
            render_dir = Path(tmp)
            (render_dir / "slide-1.png").write_bytes(b"page")
            (render_dir / "slide-artifact-tool.png").write_bytes(b"preview")
            (render_dir / "slide-not-a-page.png").write_bytes(b"noise")
            self.assertEqual(
                builder.matching_rendered_pages(render_dir, "slide"),
                [render_dir / "slide-1.png"],
            )

    def test_classification_slide_source_footer_is_projection_readable(self) -> None:
        slide_source = (ROOT / "scripts/build_aug19_classification_slide.mjs").read_text(
            encoding="utf-8"
        )
        self.assertIn('{ fontSize: 15, color: "#6A7489" }', slide_source)

    def test_hebrew_control_tokens_use_lrm_and_nonbreaking_hyphens(self) -> None:
        from docx import Document

        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_rtl_token", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        paragraph = Document().add_paragraph()
        builder.add_text_runs(paragraph, "הפניה [C2-ACL-06] והניסוי EXP-005.", rtl=True)
        xml = paragraph._p.xml
        self.assertIn("\u200eC2\u2011ACL\u201106\u200e", xml)
        self.assertIn("\u200eEXP\u2011005\u200e", xml)

    def test_generated_tables_prevent_rows_from_splitting_across_pages(self) -> None:
        from docx import Document

        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_table_rows", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        table = builder.add_table(
            Document(),
            ["כותרת", "מצב"],
            [["ייחוס ומקור ראייתי", "ממתין"]],
            rtl=True,
        )
        self.assertEqual(table._tbl.xml.count("w:cantSplit"), 2)

    def test_hebrew_paragraphs_and_bullets_use_visual_right_alignment(self) -> None:
        from docx import Document

        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_rtl_alignment", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        document = Document()
        paragraph = document.add_paragraph()
        builder.add_text_runs(paragraph, "כותרת", rtl=True)
        self.assertIn('w:jc w:val="left"', paragraph._p.xml)
        bullet = builder.add_bullet(document, "פריט", rtl=True)
        self.assertIn('w:right="540"', bullet._p.xml)
        self.assertNotIn('w:left="540"', bullet._p.xml)

    def test_non_rtl_title_block_is_left_aligned_not_justified(self) -> None:
        from docx import Document

        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_title_alignment", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        document = Document()
        builder.add_title_block(document, "English | עברית", "Status | מצב", rtl=False)
        for paragraph in document.paragraphs:
            self.assertIn('w:jc w:val="left"', paragraph._p.xml)

    def test_bilingual_tracker_cells_render_hebrew_in_a_separate_bidi_paragraph(self) -> None:
        from docx import Document

        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_bilingual_cells", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        table = builder.add_table(
            Document(),
            ["Control", "Evidence"],
            [["C", "EN: English evidence\nHE: ראיה בעברית"]],
            rtl=False,
        )
        cell = table.rows[1].cells[1]
        self.assertEqual(len(cell.paragraphs), 2)
        self.assertNotIn("w:bidi", cell.paragraphs[0]._p.xml)
        self.assertIn("w:bidi", cell.paragraphs[1]._p.xml)

    def test_hebrew_preread_starts_decisions_on_page_two(self) -> None:
        from docx import Document

        spec = importlib.util.spec_from_file_location(
            "build_aug19_supervisor_package_he_preread", BUILDER_SCRIPT
        )
        assert spec and spec.loader
        builder = importlib.util.module_from_spec(spec)
        sys.modules[spec.name] = builder
        spec.loader.exec_module(builder)
        with tempfile.TemporaryDirectory() as tmp:
            output = Path(tmp) / "preread-he.docx"
            builder.build_preread(builder.load_content(), "he", output)
            document = Document(output)
            heading = next(p for p in document.paragraphs if p.text == "החלטות נדרשות")
            self.assertIn("w:pageBreakBefore", heading._p.xml)

    def test_hebrew_chapter_pdf_uses_unambiguous_corpus_reduction_wording(self) -> None:
        reader = __import__("pypdf").PdfReader(
            str(FINAL_DIR / "Chapter_2_Literature_Review_HE.pdf")
        )
        text = "\n".join(page.extract_text() or "" for page in reader.pages)

        self.assertIn("המתכנסים", text)
        self.assertNotIn("->", text)
        self.assertNotIn("<-", text)

    def test_tracked_content_and_final_package_pass_the_release_validator(self) -> None:
        content = json.loads(CONTENT.read_text(encoding="utf-8"))
        MODULE.validate_content_contract(content)
        report = MODULE.validate_package(FINAL_DIR, content_path=CONTENT)

        self.assertEqual(report["file_count"], 10)
        self.assertEqual(report["manifest_hashes_verified"], 9)
        self.assertTrue(report["docx_pdf_page_parity"])
        self.assertEqual(report["placeholders_found"], [])
        self.assertEqual(report["forbidden_claims_found"], [])


if __name__ == "__main__":
    unittest.main()
